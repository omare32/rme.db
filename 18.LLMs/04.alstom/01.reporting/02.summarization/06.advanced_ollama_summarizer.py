import os
import json
import requests
import time
from typing import List, Dict, Any
import PyPDF2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import defaultdict

# Configuration
OLLAMA_API_URL = "http://10.10.12.202:11434/api/generate"  # GPU machine's IP
MODEL_NAME = "mistral"  # Using mistral model which is available on the server

def load_all_pdf_files():
    """Get all PDF files from the Alstom directory"""
    base_paths = ["C:\\alstom\\folder1", "C:\\alstom\\folder2"]
    pdf_files = []
    
    for base_path in base_paths:
        for root, _, files in os.walk(base_path):
            for file in files:
                if file.lower().endswith('.pdf'):
                    full_path = os.path.join(root, file)
                    simplified_path = os.path.relpath(os.path.dirname(full_path), base_path)
                    pdf_files.append({
                        'name': file,
                        'full_path': full_path,
                        'simplified_path': simplified_path
                    })
    
    print(f"Found {len(pdf_files)} PDF files in total")
    return pdf_files

def extract_pdf_text(pdf_path):
    """Extract text content from a PDF file"""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
    except Exception as e:
        print(f"Error extracting text from {pdf_path}: {str(e)}")
        return None
        
def generate_keyword_summary(text, doc_name):
    """Generate a summary based on keyword extraction (fallback method)"""
    if not text or len(text.strip()) < 100:
        return f"Error: Document '{doc_name}' has insufficient text content for summarization."
    
    # Extract key sentences based on important keywords
    important_keywords = [
        "specification", "requirement", "scope", "design", "plan",
        "schedule", "budget", "approval", "compliance", "standard",
        "regulation", "safety", "quality", "deadline", "milestone",
        "contract", "agreement", "responsibility", "deliverable"
    ]
    
    # Split text into sentences
    sentences = text.replace('\n', ' ').split('. ')
    
    # Score sentences based on keyword matches
    scored_sentences = []
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 10:  # Skip very short sentences
            continue
            
        score = 0
        for keyword in important_keywords:
            if keyword in sentence.lower():
                score += 1
                
        if score > 0:
            scored_sentences.append((sentence, score))
    
    # Sort sentences by score (highest first)
    scored_sentences.sort(key=lambda x: x[1], reverse=True)
    
    # Take top sentences (up to 10)
    top_sentences = [s[0] for s in scored_sentences[:10]]
    
    # Create summary
    if top_sentences:
        summary = f"Document Summary for '{doc_name}':\n\n"
        summary += "This document contains information about " + ", ".join(
            [kw for kw in important_keywords if any(kw in s.lower() for s in top_sentences)]
        ) + ".\n\n"
        summary += "Key points:\n"
        summary += "\n".join([f"- {s}" for s in top_sentences])
        return summary
    else:
        return f"Could not generate summary for '{doc_name}'. Document may not contain relevant keywords."

def summarize_with_ollama(text, doc_name):
    """Use Ollama API to summarize document text"""
    if not text or len(text.strip()) < 100:
        return f"Error: Document '{doc_name}' has insufficient text content for summarization."
    
    # Truncate text if it's too long (Ollama models typically have context limits)
    max_chars = 8000  # Reduced to handle Mistral's context window
    if len(text) > max_chars:
        text = text[:max_chars] + "..."
    
    prompt = f"""Summarize the following document titled '{doc_name}'. 
Focus on key information, main points, and important details. 
Keep the summary concise (200-300 words) but comprehensive.

DOCUMENT TEXT:
{text}

SUMMARY:"""
    
    # Try both API formats
    try:
        # First try the generate API
        payload = {
            "model": MODEL_NAME,
            "prompt": prompt,
            "stream": False
        }
        
        response = requests.post(OLLAMA_API_URL, json=payload)
        
        if response.status_code == 200:
            result = response.json()
            return result.get('response', '')
        
        # If that fails, try the completions API
        completions_url = OLLAMA_API_URL.replace('/generate', '/completions')
        payload = {
            "model": MODEL_NAME,
            "prompt": prompt,
            "stream": False,
            "max_tokens": 500
        }
        
        response = requests.post(completions_url, json=payload)
        
        if response.status_code == 200:
            result = response.json()
            return result.get('choices', [{}])[0].get('text', '')
            
        # If both fail, use our local keyword summarizer as fallback
        return generate_keyword_summary(text, doc_name)
        
    except Exception as e:
        print(f"Error connecting to Ollama API: {str(e)}")
        # Fallback to local keyword summarization
        return generate_keyword_summary(text, doc_name)

def categorize_with_ollama(docs_with_summaries):
    """Use Ollama to categorize documents into groups"""
    # Create a simplified list of documents with their summaries
    doc_list = []
    for doc in docs_with_summaries:
        name = doc['name']
        summary = doc.get('summary', 'No summary available')
        doc_list.append(f"Document: {name}\nSummary: {summary}\n")
    
    # Join the list with separators
    doc_text = "\n---\n".join(doc_list)
    
    # Truncate if too long
    max_chars = 8000
    if len(doc_text) > max_chars:
        # Take a representative sample of documents
        doc_text = "\n---\n".join(doc_list[:20])
    
    prompt = f"""I have a collection of project documents that need to be categorized into logical groups.
Based on the document names and summaries below, suggest 5-8 categories that would be useful for organizing these documents.
For each category, explain what types of documents should be included.

DOCUMENTS:
{doc_text}

CATEGORIES:"""
    
    try:
        # Try the generate API
        payload = {
            "model": MODEL_NAME,
            "prompt": prompt,
            "stream": False
        }
        
        response = requests.post(OLLAMA_API_URL, json=payload)
        
        if response.status_code == 200:
            result = response.json()
            categories_text = result.get('response', '')
            
            # Parse the categories
            categories = parse_categories(categories_text)
            if categories:
                return categories
        
        # If Ollama fails, use default categories
        return get_default_categories()
        
    except Exception as e:
        print(f"Error connecting to Ollama API for categorization: {str(e)}")
        return get_default_categories()

def parse_categories(categories_text):
    """Parse categories from Ollama's response"""
    categories = {}
    
    # Look for numbered or bulleted list items
    pattern = r'(?:\d+\.|\*|\-)\s+(.*?)(?::|$)'
    matches = re.findall(pattern, categories_text)
    
    if matches:
        for i, match in enumerate(matches):
            category_name = match.strip()
            categories[category_name] = []
    else:
        # Try to find category headers (capitalized words followed by colon)
        pattern = r'([A-Z][A-Za-z\s]+):'
        matches = re.findall(pattern, categories_text)
        
        if matches:
            for match in matches:
                category_name = match.strip()
                categories[category_name] = []
    
    if not categories:
        return None
    
    return categories

def get_default_categories():
    """Return default document categories if Ollama categorization fails"""
    return {
        "Technical Specifications": [],
        "Design Drawings": [],
        "Project Management": [],
        "Contracts and Agreements": [],
        "Reports and Assessments": [],
        "Safety and Compliance": [],
        "Miscellaneous": []
    }

def assign_documents_to_categories(docs_with_summaries, categories):
    """Assign documents to appropriate categories"""
    # Define keywords for each category
    category_keywords = {
        "Technical Specifications": ["specification", "spec", "technical", "requirement", "standard"],
        "Design Drawings": ["drawing", "design", "plan", "layout", "sketch", "diagram", "ifc"],
        "Project Management": ["schedule", "timeline", "milestone", "progress", "management", "plan"],
        "Contracts and Agreements": ["contract", "agreement", "legal", "terms", "condition", "scope"],
        "Reports and Assessments": ["report", "assessment", "analysis", "evaluation", "survey"],
        "Safety and Compliance": ["safety", "compliance", "regulation", "standard", "quality", "inspection"],
        "Miscellaneous": []
    }
    
    # Assign each document to a category
    categorized_docs = {category: [] for category in categories}
    
    for doc in docs_with_summaries:
        assigned = False
        doc_text = (doc['name'] + ' ' + doc.get('summary', '')).lower()
        
        # Check each category
        for category, keywords in category_keywords.items():
            if category in categories:
                for keyword in keywords:
                    if keyword in doc_text:
                        categorized_docs[category].append(doc)
                        assigned = True
                        break
                
                if assigned:
                    break
        
        # If not assigned to any category, put in Miscellaneous
        if not assigned and "Miscellaneous" in categories:
            categorized_docs["Miscellaneous"].append(doc)
    
    return categorized_docs

def save_summaries(summaries):
    """Save summaries to a JSON file"""
    data_dir = os.path.join(os.path.dirname(__file__), "data")
    os.makedirs(data_dir, exist_ok=True)
    
    output_path = os.path.join(data_dir, "ollama_summaries.json")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(summaries, f, indent=2, ensure_ascii=False)
    
    print(f"Saved summaries to {output_path}")
    return output_path

def load_summaries():
    """Load existing summaries if available"""
    data_dir = os.path.join(os.path.dirname(__file__), "data")
    summaries_path = os.path.join(data_dir, "ollama_summaries.json")
    
    if not os.path.exists(summaries_path):
        return {}
    
    try:
        with open(summaries_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return {}

def create_category_report(category_name, docs, output_path):
    """Create a Word document with summaries for a specific category"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Alstom Project Documentation: {category_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    intro = doc.add_paragraph()
    intro.add_run('Executive Summary').bold = True
    doc.add_paragraph(f'This report contains summaries of important {category_name.lower()} documents from the Alstom project. These documents have been selected based on their relevance to key project aspects and organized to provide a comprehensive overview.')
    
    doc.add_paragraph()  # Add space
    
    # Add document summaries
    doc.add_heading('Document Summaries', 1)
    
    for i, doc_info in enumerate(docs, 1):
        name = doc_info['name']
        path = doc_info['simplified_path']
        
        # Add document header with formatting
        header = doc.add_paragraph()
        header.add_run(f"{i}. {name}").bold = True
        header.add_run(f"\nLocation: {path}")
        
        # Add summary with formatting
        if 'summary' in doc_info:
            summary = doc_info['summary']
            p = doc.add_paragraph()
            p.add_run('Summary: ').bold = True
            p.add_run(summary)
        else:
            doc.add_paragraph('Summary not available for this document.').italic = True
            
        doc.add_paragraph()  # Add space between documents
    
    # Save the document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\nReport saved to: {output_path}")

def create_master_report(categorized_docs, output_path):
    """Create a master Word document with all documents organized by category"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Alstom Project Documentation Master Index', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    intro = doc.add_paragraph()
    intro.add_run('Executive Summary').bold = True
    doc.add_paragraph('This report provides a comprehensive index of all important documents from the Alstom project, organized by category. Each category has its own detailed report with full summaries.')
    
    doc.add_paragraph()  # Add space
    
    # Add table of contents header
    doc.add_heading('Document Categories', 1)
    
    # Add each category
    for category, docs in categorized_docs.items():
        if docs:  # Only add category if it has documents
            category_heading = doc.add_heading(category, 2)
            
            # Add document list
            for i, doc_info in enumerate(docs, 1):
                name = doc_info['name']
                path = doc_info['simplified_path']
                
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{name}").bold = True
                p.add_run(f" (in {path})")
            
            doc.add_paragraph()  # Add space between categories
    
    # Save the document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\nMaster index saved to: {output_path}")

def main():
    # Load all PDF files
    pdf_files = load_all_pdf_files()
    
    # Load existing summaries
    summaries = load_summaries()
    print(f"Loaded {len(summaries)} existing summaries.")
    
    # Process documents
    docs_with_summaries = []
    for i, doc in enumerate(pdf_files, 1):
        name = doc['name']
        
        if name in summaries:
            print(f"Skipping {i}/{len(pdf_files)}: {name} (already summarized)")
            # Add to docs_with_summaries
            doc_with_summary = doc.copy()
            doc_with_summary['summary'] = summaries[name]
            docs_with_summaries.append(doc_with_summary)
            continue
        
        print(f"\nProcessing {i}/{len(pdf_files)}: {name}")
        text = extract_pdf_text(doc['full_path'])
        
        if text:
            print(f"Extracted {len(text)} characters of text. Summarizing...")
            summary = summarize_with_ollama(text, name)
            summaries[name] = summary
            
            # Add to docs_with_summaries
            doc_with_summary = doc.copy()
            doc_with_summary['summary'] = summary
            docs_with_summaries.append(doc_with_summary)
            
            # Save progress after each summary
            save_summaries(summaries)
        else:
            print(f"Warning: Could not extract text from {name}")
            summaries[name] = "Error: Could not read document"
            
            # Still add to docs_with_summaries
            doc_with_summary = doc.copy()
            doc_with_summary['summary'] = "Error: Could not read document"
            docs_with_summaries.append(doc_with_summary)
    
    # Categorize documents
    print("\nCategorizing documents...")
    categories = categorize_with_ollama(docs_with_summaries)
    categorized_docs = assign_documents_to_categories(docs_with_summaries, categories)
    
    # Create reports directory
    reports_dir = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(reports_dir, exist_ok=True)
    
    # Create category reports
    print("\nGenerating category reports...")
    for category, docs in categorized_docs.items():
        if docs:  # Only create report if category has documents
            category_filename = category.replace(" ", "_").lower()
            output_path = os.path.join(reports_dir, f"Alstom_{category_filename}_ollama.docx")
            create_category_report(category, docs, output_path)
    
    # Create master report
    master_path = os.path.join(reports_dir, "Alstom_Master_Index_ollama.docx")
    create_master_report(categorized_docs, master_path)
    
    print("\nAll reports generated successfully!")

if __name__ == "__main__":
    main()
