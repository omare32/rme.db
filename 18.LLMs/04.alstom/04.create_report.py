import os
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_data():
    """Load existing summaries and document info"""
    data_dir = os.path.join(os.path.dirname(__file__), "data")
    
    # Try different encodings
    encodings = ['utf-8', 'utf-16', 'latin1', 'ascii']
    
    # Load important docs
    docs_path = os.path.join(data_dir, "important_docs.json")
    if not os.path.exists(docs_path):
        print(f"No document list found at {docs_path}")
        return None, None
    
    important_docs = None
    for encoding in encodings:
        try:
            with open(docs_path, 'r', encoding=encoding) as f:
                important_docs = json.load(f)
            break
        except (UnicodeDecodeError, json.JSONDecodeError):
            continue
    
    if not important_docs:
        print(f"Could not read {docs_path} with any encoding")
        return None, None
    
    # Load summaries
    summaries_path = os.path.join(data_dir, "summaries.json")
    if not os.path.exists(summaries_path):
        print(f"No summaries found at {summaries_path}")
        return None, None
    
    summaries = None
    for encoding in encodings:
        try:
            with open(summaries_path, 'r', encoding=encoding) as f:
                summaries = json.load(f)
            break
        except (UnicodeDecodeError, json.JSONDecodeError):
            continue
    
    if not summaries:
        print(f"Could not read {summaries_path} with any encoding")
        return None, None
    
    return important_docs, summaries

def create_report(important_docs, summaries, output_path):
    """Create a nicely formatted Word document with summaries"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Alstom Project Documentation Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    intro = doc.add_paragraph()
    intro.add_run('Executive Summary').bold = True
    doc.add_paragraph('This report contains summaries of the most important documents from the Alstom project. These documents have been selected based on their relevance to key project aspects such as specifications, technical requirements, plans, and critical correspondence.')
    
    doc.add_paragraph()  # Add space
    
    # Add document summaries
    doc.add_heading('Document Summaries', 1)
    
    # Group documents by type
    doc_groups = {
        'Specifications': [],
        'Drawings': [],
        'Reports': [],
        'Plans': [],
        'Other': []
    }
    
    for doc_info in important_docs:
        name = doc_info['name'].lower()
        if 'specification' in name:
            doc_groups['Specifications'].append(doc_info)
        elif 'drawing' in name or 'ifc' in name:
            doc_groups['Drawings'].append(doc_info)
        elif 'report' in name:
            doc_groups['Reports'].append(doc_info)
        elif 'plan' in name:
            doc_groups['Plans'].append(doc_info)
        else:
            doc_groups['Other'].append(doc_info)
    
    # Add documents by group
    doc_num = 1
    for group_name, group_docs in doc_groups.items():
        if group_docs:  # Only add group if it has documents
            doc.add_heading(group_name, 2)
            
            for doc_info in group_docs:
                name = doc_info['name']
                path = doc_info['simplified_path']
                
                # Add document header with formatting
                header = doc.add_paragraph()
                header.add_run(f"{doc_num}. {name}").bold = True
                header.add_run(f"\nLocation: {path}")
                
                # Add summary with formatting
                if name in summaries:
                    summary = summaries[name]
                    p = doc.add_paragraph()
                    p.add_run('Summary: ').bold = True
                    p.add_run(summary)
                else:
                    doc.add_paragraph('Summary not available for this document.').italic = True
                    
                doc.add_paragraph()  # Add space between documents
                doc_num += 1
    
    # Save the document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\nReport saved to: {output_path}")

def main():
    print("Loading existing document data...")
    important_docs, summaries = load_data()
    
    if not important_docs or not summaries:
        print("Error: Could not find required data files.")
        return
    
    print(f"\nFound {len(important_docs)} documents with {len(summaries)} summaries.")
    
    # Create report
    output_dir = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Alstom_Project_Documentation_Summary.docx")
    
    print("\nCreating report...")
    try:
        create_report(important_docs, summaries, output_path)
    except Exception as e:
        print(f"Error creating report: {str(e)}")
        return

if __name__ == "__main__":
    main()
