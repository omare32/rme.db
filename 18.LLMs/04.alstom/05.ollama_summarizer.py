import os
import json
import requests
import time
from typing import List, Dict, Any
import PyPDF2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
OLLAMA_API_URL = "http://10.10.12.202:11434/api/generate"  # GPU machine's IP
MODEL_NAME = "mistral"  # Using mistral model which is available on the server

def load_data():
    """Load existing document data"""
    data_dir = os.path.join(os.path.dirname(__file__), "data")
    
    # Try different encodings
    encodings = ['utf-8', 'utf-16', 'latin1', 'ascii']
    
    # Load important docs
    docs_path = os.path.join(data_dir, "important_docs.json")
    if not os.path.exists(docs_path):
        print(f"No document list found at {docs_path}")
        return None
    
    important_docs = None
    for encoding in encodings:
        try:
            with open(docs_path, 'r', encoding=encoding) as f:
                important_docs = json.load(f)
            break
        except (UnicodeDecodeError, json.JSONDecodeError):
            continue
    
    if not important_docs:
        print(f"Could not read {docs_path} with any encoding")
        return None
    
    return important_docs

def extract_pdf_text(pdf_path):
    """Extract text content from a PDF file"""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
    except Exception as e:
        print(f"Error extracting text from {pdf_path}: {str(e)}")
        return None
        
def generate_keyword_summary(text, doc_name):
    """Generate a summary based on keyword extraction (fallback method)"""
    if not text or len(text.strip()) < 100:
        return f"Error: Document '{doc_name}' has insufficient text content for summarization."
    
    # Extract key sentences based on important keywords
    important_keywords = [
        "specification", "requirement", "scope", "design", "plan",
        "schedule", "budget", "approval", "compliance", "standard",
        "regulation", "safety", "quality", "deadline", "milestone",
        "contract", "agreement", "responsibility", "deliverable"
    ]
    
    # Split text into sentences
    sentences = text.replace('\n', ' ').split('. ')
    
    # Score sentences based on keyword matches
    scored_sentences = []
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 10:  # Skip very short sentences
            continue
            
        score = 0
        for keyword in important_keywords:
            if keyword in sentence.lower():
                score += 1
                
        if score > 0:
            scored_sentences.append((sentence, score))
    
    # Sort sentences by score (highest first)
    scored_sentences.sort(key=lambda x: x[1], reverse=True)
    
    # Take top sentences (up to 10)
    top_sentences = [s[0] for s in scored_sentences[:10]]
    
    # Create summary
    if top_sentences:
        summary = f"Document Summary for '{doc_name}':\n\n"
        summary += "This document contains information about " + ", ".join(
            [kw for kw in important_keywords if any(kw in s.lower() for s in top_sentences)]
        ) + ".\n\n"
        summary += "Key points:\n"
        summary += "\n".join([f"- {s}" for s in top_sentences])
        return summary
    else:
        return f"Could not generate summary for '{doc_name}'. Document may not contain relevant keywords."

def summarize_with_ollama(text, doc_name):
    """Use Ollama API to summarize document text"""
    if not text or len(text.strip()) < 100:
        return f"Error: Document '{doc_name}' has insufficient text content for summarization."
    
    # Truncate text if it's too long (Ollama models typically have context limits)
    max_chars = 8000  # Reduced to handle Mistral's context window
    if len(text) > max_chars:
        text = text[:max_chars] + "..."
    
    prompt = f"""Summarize the following document titled '{doc_name}'. 
Focus on key information, main points, and important details. 
Keep the summary concise (200-300 words) but comprehensive.

DOCUMENT TEXT:
{text}

SUMMARY:"""
    
    # Try both API formats
    try:
        # First try the generate API
        payload = {
            "model": MODEL_NAME,
            "prompt": prompt,
            "stream": False
        }
        
        response = requests.post(OLLAMA_API_URL, json=payload)
        
        if response.status_code == 200:
            result = response.json()
            return result.get('response', '')
        
        # If that fails, try the completions API
        completions_url = OLLAMA_API_URL.replace('/generate', '/completions')
        payload = {
            "model": MODEL_NAME,
            "prompt": prompt,
            "stream": False,
            "max_tokens": 500
        }
        
        response = requests.post(completions_url, json=payload)
        
        if response.status_code == 200:
            result = response.json()
            return result.get('choices', [{}])[0].get('text', '')
            
        # If both fail, use our local keyword summarizer as fallback
        return generate_keyword_summary(text, doc_name)
        
    except Exception as e:
        print(f"Error connecting to Ollama API: {str(e)}")
        # Fallback to local keyword summarization
        return generate_keyword_summary(text, doc_name)

def save_summaries(summaries):
    """Save summaries to a JSON file"""
    data_dir = os.path.join(os.path.dirname(__file__), "data")
    os.makedirs(data_dir, exist_ok=True)
    
    output_path = os.path.join(data_dir, "summaries.json")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(summaries, f, indent=2, ensure_ascii=False)
    
    print(f"Saved summaries to {output_path}")
    return output_path

def create_report(important_docs, summaries, output_path):
    """Create a Word document with summaries"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Alstom Project Documentation Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    intro = doc.add_paragraph()
    intro.add_run('Executive Summary').bold = True
    doc.add_paragraph('This report contains summaries of the most important documents from the Alstom project. These documents have been selected based on their relevance to key project aspects such as specifications, technical requirements, plans, and critical correspondence.')
    
    doc.add_paragraph()  # Add space
    
    # Add document summaries
    doc.add_heading('Document Summaries', 1)
    
    # Group documents by type
    doc_groups = {
        'Specifications': [],
        'Drawings': [],
        'Reports': [],
        'Plans': [],
        'Other': []
    }
    
    for doc_info in important_docs:
        name = doc_info['name'].lower()
        if 'specification' in name:
            doc_groups['Specifications'].append(doc_info)
        elif 'drawing' in name or 'ifc' in name:
            doc_groups['Drawings'].append(doc_info)
        elif 'report' in name:
            doc_groups['Reports'].append(doc_info)
        elif 'plan' in name:
            doc_groups['Plans'].append(doc_info)
        else:
            doc_groups['Other'].append(doc_info)
    
    # Add documents by group
    doc_num = 1
    for group_name, group_docs in doc_groups.items():
        if group_docs:  # Only add group if it has documents
            doc.add_heading(group_name, 2)
            
            for doc_info in group_docs:
                name = doc_info['name']
                path = doc_info['simplified_path']
                
                # Add document header with formatting
                header = doc.add_paragraph()
                header.add_run(f"{doc_num}. {name}").bold = True
                header.add_run(f"\nLocation: {path}")
                
                # Add summary with formatting
                if name in summaries:
                    summary = summaries[name]
                    p = doc.add_paragraph()
                    p.add_run('Summary: ').bold = True
                    p.add_run(summary)
                else:
                    doc.add_paragraph('Summary not available for this document.').italic = True
                    
                doc.add_paragraph()  # Add space between documents
                doc_num += 1
    
    # Save the document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\nReport saved to: {output_path}")

def main():
    print("Loading document data...")
    important_docs = load_data()
    
    if not important_docs:
        print("Error: Could not find or load document data.")
        return
    
    print(f"\nFound {len(important_docs)} documents to process.")
    
    # Load existing summaries if available
    data_dir = os.path.join(os.path.dirname(__file__), "data")
    summaries_path = os.path.join(data_dir, "summaries.json")
    
    try:
        with open(summaries_path, 'r', encoding='utf-8') as f:
            summaries = json.load(f)
        print(f"Loaded {len(summaries)} existing summaries.")
    except:
        summaries = {}
        print("No existing summaries found or could not load them.")
    
    # Process documents
    for i, doc in enumerate(important_docs, 1):
        name = doc['name']
        
        if name in summaries:
            print(f"Skipping {i}/{len(important_docs)}: {name} (already summarized)")
            continue
        
        print(f"\nProcessing {i}/{len(important_docs)}: {name}")
        text = extract_pdf_text(doc['full_path'])
        
        if text:
            print(f"Extracted {len(text)} characters of text. Summarizing...")
            summary = summarize_with_ollama(text, name)
            summaries[name] = summary
            
            # Save progress after each summary
            save_summaries(summaries)
        else:
            print(f"Warning: Could not extract text from {name}")
            summaries[name] = "Error: Could not read document"
    
    # Create report
    print("\nGenerating summary report...")
    output_dir = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Alstom_Project_Documentation_Summary.docx")
    
    try:
        create_report(important_docs, summaries, output_path)
    except Exception as e:
        print(f"Error creating report: {str(e)}")

if __name__ == "__main__":
    main()
