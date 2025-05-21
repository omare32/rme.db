import os
import sys
import json
import PyPDF2
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
from dotenv import load_dotenv
from pathlib import Path
import time
from typing import List, Dict, Any

def load_openai_key():
    """Load OpenAI API key from .env file"""
    env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
    
    if not os.path.exists(env_path):
        raise ValueError(".env file not found")
    
    # Try different encodings
    encodings = ['utf-8', 'utf-16', 'latin1', 'ascii']
    key = None
    
    for encoding in encodings:
        try:
            with open(env_path, 'r', encoding=encoding) as f:
                for line in f:
                    if line.startswith('OPENAI_API_KEY='):
                        key = line.strip().split('=', 1)[1].strip().strip('"').strip("'")
                        break
            if key:
                break
        except UnicodeDecodeError:
            continue
    
    if not key:
        raise ValueError("OpenAI API key not found in .env file")
    
    openai.api_key = key

def get_pdf_files(base_path):
    """Extract PDF files with simplified paths (last 2 folders)"""
    pdf_files = []
    
    for root, _, files in os.walk(base_path):
        for file in files:
            if file.lower().endswith('.pdf'):
                # Get full path
                full_path = os.path.join(root, file)
                # Get relative path from base_path
                rel_path = os.path.relpath(root, base_path)
                # Split path and take last 2 parts if available
                path_parts = rel_path.split(os.sep)
                simplified_path = os.path.join(*path_parts[-2:]) if len(path_parts) >= 2 else rel_path
                
                pdf_files.append({
                    'name': file,
                    'simplified_path': simplified_path,
                    'full_path': full_path
                })
    
    return pdf_files

def call_openai_with_retry(messages: List[Dict[str, str]], max_retries: int = 3, initial_wait: int = 20) -> str:
    """Call OpenAI API with exponential backoff retry"""
    wait_time = initial_wait
    
    for attempt in range(max_retries):
        try:
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=messages,
                temperature=0.3
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            if attempt == max_retries - 1:
                raise e
            print(f"Attempt {attempt + 1} failed: {str(e)}. Waiting {wait_time} seconds...")
            time.sleep(wait_time)
            wait_time *= 2

def keyword_filter_documents(pdf_files: List[Dict[str, str]], max_docs: int = 50) -> List[Dict[str, str]]:
    """Filter documents based on important keywords in their names and paths"""
    # Define keywords and their importance weights
    keywords = {
        'scope': 10,
        'specification': 9,
        'spec': 9,
        'drawing': 8,
        'plan': 8,
        'approval': 7,
        'approved': 7,
        'design': 7,
        'technical': 6,
        'requirement': 6,
        'method': 6,
        'statement': 6,
        'contract': 5,
        'agreement': 5,
        'schedule': 5,
        'timeline': 5,
        'budget': 5,
        'cost': 5,
        'quality': 4,
        'safety': 4,
        'report': 3,
        'review': 3,
        'final': 3
    }
    
    # Score each document
    scored_docs = []
    for doc in pdf_files:
        score = 0
        # Combine name and path for searching
        text = (doc['name'] + ' ' + doc['simplified_path']).lower()
        
        # Calculate score based on keyword matches
        for keyword, weight in keywords.items():
            if keyword in text:
                score += weight
                # Bonus for keywords in the filename
                if keyword in doc['name'].lower():
                    score += weight
        
        if score > 0:  # Only include documents that matched at least one keyword
            scored_docs.append((score, doc))
    
    # Sort by score and return top documents
    scored_docs.sort(key=lambda x: x[0], reverse=True)
    return [doc for _, doc in scored_docs[:max_docs]]

def get_important_documents(pdf_files: List[Dict[str, str]]) -> List[Dict[str, str]]:
    """Get the most important documents from the list of PDF files"""
    print("\nFiltering documents based on keywords...")
    filtered_docs = keyword_filter_documents(pdf_files)
    print(f"Found {len(filtered_docs)} potentially important documents")
    
    if not filtered_docs:
        print("No documents found matching important keywords")
        return None
    
    # Since we already know which documents are important from our previous work,
    # let's use the specifications and key documents we found
    important_patterns = [
        "Specification",
        "IFC",
        "Report",
        "Contract",
        "Plan",
        "Matrix"
    ]
    
    selected_docs = []
    for doc in filtered_docs:
        name = doc['name']
        if any(pattern.lower() in name.lower() for pattern in important_patterns):
            selected_docs.append(doc)
            if len(selected_docs) >= 20:
                break
    
    if selected_docs:
        print("\nSelected documents:")
        for i, doc in enumerate(selected_docs, 1):
            print(f"{i}. {doc['name']} (in {doc['simplified_path']})")
        return selected_docs
    
    print("Could not identify important documents")
    return None

def extract_pdf_text(pdf_path):
    """Extract text content from a PDF file"""
    try:
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {str(e)}")
        return None

def summarize_document(text: str, filename: str) -> str:
    """Use OpenAI to generate a summary of the document"""
    if not text:
        return f"Error: Could not read {filename}"
    
    prompt = f"""Summarize the following document from an Alstom Factory construction project. Focus on key points that would be important for a new project manager. Keep the summary concise (max 200 words).

Document name: {filename}

Content:
{text[:4000]}  # Limiting content to avoid token limits
"""

    try:
        messages = [
            {"role": "system", "content": "You are a helpful assistant that summarizes technical documents."},
            {"role": "user", "content": prompt}
        ]
        
        return call_openai_with_retry(messages)
        
    except Exception as e:
        return f"Error summarizing {filename}: {str(e)}"

def create_summary_report(summaries, output_path):
    """Generate a Word document with all summaries"""
    doc = Document()
    
    # Title
    doc.add_heading('Key Project Documents Summary', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('This report contains summaries of the 20 most important documents identified for the Alstom Factory construction project. These documents have been selected and summarized to provide a quick overview of critical project information for the new project manager.')
    
    # Document Summaries
    doc.add_heading('Document Summaries', level=1)
    for i, summary in enumerate(summaries, 1):
        # Document header
        doc.add_heading(f'{i}. {summary["name"]}', level=2)
        doc.add_paragraph(f'Location: {summary["path"]}')
        
        # Summary content
        doc.add_paragraph(summary["content"])
        
        # Add spacing between summaries
        doc.add_paragraph()
    
    # Save the document
    doc.save(output_path)
    return output_path

def save_progress(data, filename):
    """Save progress to a JSON file"""
    output_dir = os.path.join(os.path.dirname(__file__), "data")
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, filename)
    
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"Saved {len(data) if isinstance(data, list) else len(data.keys())} items to {filepath}")
    return filepath

def fix_corrupted_json(filepath):
    """Try to fix corrupted JSON file by reading it line by line"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
            return json.loads(content)
    except json.JSONDecodeError:
        print(f"Attempting to fix corrupted JSON in {filepath}...")
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                # Try to parse as much as we can
                data = {}
                current_key = None
                current_value = []
                
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                        
                    if line.startswith('"') and ':' in line:
                        # Save previous entry if exists
                        if current_key and current_value:
                            data[current_key] = '\n'.join(current_value)
                            current_value = []
                        
                        # Start new entry
                        parts = line.split('":', 1)
                        if len(parts) == 2:
                            current_key = parts[0].strip('" ')
                            value_part = parts[1].strip().strip('"')
                            if value_part:
                                current_value.append(value_part)
                    elif current_key and line:
                        current_value.append(line.strip('" '))
                
                # Save last entry
                if current_key and current_value:
                    data[current_key] = '\n'.join(current_value)
                
                return data
        except Exception as e:
            print(f"Could not fix JSON: {str(e)}")
            return {}

def load_progress(filename):
    """Load progress from a JSON file"""
    filepath = os.path.join(os.path.dirname(__file__), "data", filename)
    if os.path.exists(filepath):
        print(f"Loading progress from {filepath}...")
        try:
            with open(filepath, encoding='utf-8') as f:
                data = json.load(f)
        except json.JSONDecodeError:
            print("JSON file is corrupted, attempting to fix...")
            data = fix_corrupted_json(filepath)
            if data:
                # Backup corrupted file
                backup_path = filepath + '.bak'
                import shutil
                shutil.copy2(filepath, backup_path)
                # Save fixed version
                with open(filepath, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2, ensure_ascii=False)
                print(f"Fixed JSON file saved, original backed up to {backup_path}")
        
        print(f"Loaded {len(data) if isinstance(data, list) else len(data.keys())} items")
        return data
    return None

def extract_from_word_doc(docx_path):
    """Extract document names and summaries from existing Word document"""
    if not os.path.exists(docx_path):
        print(f"Word document not found: {docx_path}")
        return None, None
        
    print(f"Extracting previous progress from {docx_path}...")
    doc = Document(docx_path)
    important_docs = []
    summaries = {}
    current_doc = None
    summary_text = []
    
    print("\nDocument paragraphs:")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"[{i}] {text[:100]}...")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Look for document headers (they start with a number followed by dot)
        if text[0].isdigit() and ". " in text[:10]:
            # Save previous summary if exists
            if current_doc and summary_text:
                summaries[current_doc['name']] = "\n".join(summary_text)
                summary_text = []
            
            # Parse new document entry
            try:
                # First split by dot and space
                _, rest = text.split(". ", 1)
                # Then try to split into name and path
                if " (in " in rest:
                    name, path = rest.split(" (in ", 1)
                    path = path.rstrip(")")
                else:
                    # If no path info, just use the whole thing as name
                    name = rest
                    path = ""
                    
                current_doc = {
                    'name': name,
                    'simplified_path': path,
                    'full_path': os.path.join("C:\\", "alstom", path, name) if path else os.path.join("C:\\", "alstom", name)
                }
                important_docs.append(current_doc)
                print(f"Found document: {name}")
            except Exception as e:
                print(f"Error parsing line: {text}")
                print(f"Error: {str(e)}")
                continue
        elif current_doc and text:
            summary_text.append(text)
    
    # Save last summary
    if current_doc and summary_text:
        summaries[current_doc['name']] = "\n".join(summary_text)
    
    if important_docs:
        print(f"Found {len(important_docs)} documents with summaries")
        return important_docs, summaries
    else:
        print("No documents found in Word file")
    
    return None, None

def create_word_report(important_docs, summaries, output_path):
    """Create a Word document with document summaries"""
    doc = Document()
    doc.add_heading('Alstom Project Documentation Analysis', 0)
    
    doc.add_paragraph('This report contains summaries of the most important documents from the Alstom project.')
    doc.add_paragraph()
    
    for i, doc_info in enumerate(important_docs, 1):
        name = doc_info['name']
        path = doc_info['simplified_path']
        
        # Add document header
        doc.add_paragraph(f"{i}. {name} (in {path})")
        
        # Add summary
        if name in summaries:
            summary = summaries[name]
            doc.add_paragraph(summary)
        else:
            doc.add_paragraph("Summary not yet available")
            
        doc.add_paragraph()  # Add space between documents
    
    # Save the document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

def get_pdf_paths():
    """Get list of PDF files from Alstom directory"""
    base_paths = ["C:\\alstom\\folder1", "C:\\alstom\\folder2"]
    pdf_files = []
    
    for base_path in base_paths:
        for root, _, files in os.walk(base_path):
            for file in files:
                if file.lower().endswith('.pdf'):
                    full_path = os.path.join(root, file)
                    simplified_path = os.path.relpath(os.path.dirname(full_path), base_path)
                    pdf_files.append({
                        'name': file,
                        'full_path': full_path,
                        'simplified_path': simplified_path
                    })
    
    return pdf_files

def main():
    try:
        print("Loading OpenAI API key...")
        load_openai_key()
        
        # Create data directory if it doesn't exist
        data_dir = os.path.join(os.path.dirname(__file__), "data")
        os.makedirs(data_dir, exist_ok=True)
        
        # Try to load from JSON files first
        important_docs = load_progress("important_docs.json")
        summaries = load_progress("summaries.json") or {}
        
        if not important_docs:
            print("\nAnalyzing PDF documents...")
            pdf_files = get_pdf_paths()
            print(f"Found {len(pdf_files)} PDF files")
            
            print("\nIdentifying important documents...")
            important_docs = get_important_documents(pdf_files)
            
            if important_docs:
                save_path = save_progress(important_docs, "important_docs.json")
                print(f"Saved selected documents to {save_path}")

        if not important_docs:
            print("No important documents were identified. Please check the file paths and try again.")
            return

        if not important_docs:
            print("No important documents were identified. Please check the file paths and try again.")
            return

        print(f"\nProcessing {len(important_docs)} documents:")
        for i, doc in enumerate(important_docs, 1):
            if doc['name'] in summaries:
                print(f"Skipping {i}/{len(important_docs)}: {doc['name']} (already processed)")
                continue
                
            print(f"\nProcessing {i}/{len(important_docs)}: {doc['name']}")
            text = extract_pdf_text(doc['full_path'])
            if text:
                print("Extracting summary...")
                summary = summarize_document(text, doc['name'])
                summaries[doc['name']] = summary
                # Save progress after each summary
                save_path = save_progress(summaries, "summaries.json")
                print(f"Saved progress to {save_path}")
            else:
                print(f"Warning: Could not extract text from {doc['name']}")
                summaries[doc['name']] = "Error: Could not read document"

        print("\nGenerating summary report...")
        output_dir = os.path.join(os.path.dirname(__file__), "reports")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, "Top_20_Documents_Summary.docx")
        
        create_word_report(important_docs, summaries, output_path)
        print(f"\nReport generated successfully: {output_path}")
        
    except Exception as e:
        print(f"\nError: {str(e)}")

if __name__ == "__main__":
    main()
