import mysql.connector
from mysql.connector import Error
import datetime
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import numpy as np

# Database configuration
DB_CONFIG = {
    'host': '10.10.11.242',
    'user': 'omar2',
    'password': 'Omar_54321',
    'database': 'RME_TEST'
}

def connect_to_database():
    """Establish connection to MySQL database"""
    try:
        connection = mysql.connector.connect(**DB_CONFIG)
        print("Connected to MySQL database")
        return connection
    except Error as e:
        print(f"Error connecting to MySQL database: {e}")
        return None

def get_enhanced_stats(connection):
    """Get enhanced statistics including project and document type information"""
    cursor = connection.cursor()
    stats = {}
    
    try:
        # Get total PDFs count
        cursor.execute("SELECT COUNT(*) FROM `po.pdfs`")
        stats['total_pdfs'] = cursor.fetchone()[0]
        
        # Get project statistics
        cursor.execute("SELECT COUNT(DISTINCT project_name) FROM `po.pdfs` WHERE project_name IS NOT NULL")
        stats['total_projects'] = cursor.fetchone()[0]
        
        # Get document type statistics
        cursor.execute("SELECT COUNT(DISTINCT document_type) FROM `po.pdfs` WHERE document_type IS NOT NULL")
        stats['total_doc_types'] = cursor.fetchone()[0]
        
        # Get document type distribution
        cursor.execute("""
        SELECT document_type, COUNT(*) as count
        FROM `po.pdfs`
        WHERE document_type IS NOT NULL
        GROUP BY document_type
        ORDER BY count DESC
        """)
        stats['doc_type_dist'] = cursor.fetchall()
        
        # Get top 10 projects by document count
        cursor.execute("""
        SELECT project_name, COUNT(*) as count
        FROM `po.pdfs`
        WHERE project_name IS NOT NULL
        GROUP BY project_name
        ORDER BY count DESC
        LIMIT 10
        """)
        stats['top_projects'] = cursor.fetchall()
        
        # Get top 5 projects first
        cursor.execute("""
        SELECT project_name, COUNT(*) as count
        FROM `po.pdfs`
        GROUP BY project_name
        ORDER BY count DESC
        LIMIT 5
        """)
        top_projects = [row[0] for row in cursor.fetchall()]
        
        # Get project-document type matrix for these projects
        project_doc_matrix = []
        for project in top_projects:
            cursor.execute("""
            SELECT 
                project_name,
                document_type,
                COUNT(*) as count
            FROM `po.pdfs`
            WHERE project_name = %s
            AND document_type IS NOT NULL
            GROUP BY project_name, document_type
            ORDER BY count DESC
            """, (project,))
            project_doc_matrix.extend(cursor.fetchall())
        
        stats['project_doc_matrix'] = project_doc_matrix
        
        # Get processing timeline
        cursor.execute("""
        SELECT 
            DATE(processed_timestamp) as process_date,
            COUNT(*) as count
        FROM `po.pdfs`
        GROUP BY process_date
        ORDER BY process_date
        """)
        stats['processing_timeline'] = cursor.fetchall()
        
    except Error as e:
        print(f"Error getting enhanced statistics: {e}")
    finally:
        cursor.close()
    
    return stats

def create_enhanced_charts(stats):
    """Create enhanced charts including project and document type visualizations"""
    charts = {}
    
    # Create folder for charts if it doesn't exist
    charts_dir = 'report_charts'
    if not os.path.exists(charts_dir):
        os.makedirs(charts_dir)
    
    # Create document type distribution chart
    if 'doc_type_dist' in stats and stats['doc_type_dist']:
        plt.figure(figsize=(10, 6))
        doc_types = [dt[0] for dt in stats['doc_type_dist']]
        counts = [dt[1] for dt in stats['doc_type_dist']]
        plt.bar(doc_types, counts)
        plt.xticks(rotation=45, ha='right')
        plt.title('Document Type Distribution')
        plt.tight_layout()
        chart_path = os.path.join(charts_dir, 'doc_type_dist.png')
        plt.savefig(chart_path)
        plt.close()
        charts['doc_type_dist'] = chart_path
    
    # Create top projects chart
    if 'top_projects' in stats and stats['top_projects']:
        plt.figure(figsize=(12, 6))
        projects = [p[0][:30] + '...' if len(p[0]) > 30 else p[0] for p in stats['top_projects']]
        counts = [p[1] for p in stats['top_projects']]
        plt.bar(projects, counts)
        plt.xticks(rotation=45, ha='right')
        plt.title('Top 10 Projects by Document Count')
        plt.tight_layout()
        chart_path = os.path.join(charts_dir, 'top_projects.png')
        plt.savefig(chart_path)
        plt.close()
        charts['top_projects'] = chart_path
    
    # Create processing timeline chart
    if 'processing_timeline' in stats and stats['processing_timeline']:
        plt.figure(figsize=(12, 6))
        dates = [pt[0] for pt in stats['processing_timeline']]
        counts = [pt[1] for pt in stats['processing_timeline']]
        plt.plot(dates, counts, marker='o')
        plt.xticks(rotation=45, ha='right')
        plt.title('PDF Processing Timeline')
        plt.xlabel('Date')
        plt.ylabel('Number of PDFs Processed')
        plt.tight_layout()
        chart_path = os.path.join(charts_dir, 'processing_timeline.png')
        plt.savefig(chart_path)
        plt.close()
        charts['processing_timeline'] = chart_path
    
    return charts

def generate_enhanced_report(stats, charts):
    """Generate an enhanced Word document report with project and document type statistics"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Enhanced PDF Extraction Database Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_paragraph.add_run(f"Report generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_run.italic = True
    
    doc.add_paragraph()
    
    # Overall Statistics section
    doc.add_heading('Overall Statistics', level=1)
    doc.add_paragraph(f"Total PDFs Processed: {stats.get('total_pdfs', 'N/A')} documents")
    doc.add_paragraph(f"Total Projects: {stats.get('total_projects', 'N/A')}")
    doc.add_paragraph(f"Total Document Types: {stats.get('total_doc_types', 'N/A')}")
    
    # Document Types section
    doc.add_heading('Document Type Analysis', level=1)
    if 'doc_type_dist' in stats:
        doc.add_paragraph("Distribution of document types:")
        for doc_type, count in stats['doc_type_dist']:
            doc.add_paragraph(f"• {doc_type}: {count} documents", style='List Bullet')
    
    if 'doc_type_dist' in charts:
        doc.add_picture(charts['doc_type_dist'], width=Inches(6))
        doc.add_paragraph()
    
    # Projects section
    doc.add_heading('Project Analysis', level=1)
    if 'top_projects' in stats:
        doc.add_paragraph("Top 10 projects by document count:")
        for project, count in stats['top_projects']:
            doc.add_paragraph(f"• {project}: {count} documents", style='List Bullet')
    
    if 'top_projects' in charts:
        doc.add_picture(charts['top_projects'], width=Inches(6))
        doc.add_paragraph()
    
    # Processing Timeline section
    doc.add_heading('Processing Timeline', level=1)
    if 'processing_timeline' in charts:
        doc.add_picture(charts['processing_timeline'], width=Inches(6))
        doc.add_paragraph()
    
    # Project-Document Type Matrix section
    if 'project_doc_matrix' in stats:
        doc.add_heading('Project-Document Type Matrix (Top 5 Projects)', level=1)
        current_project = None
        for project, doc_type, count in stats['project_doc_matrix']:
            if project != current_project:
                doc.add_paragraph(f"\n{project}:", style='Heading 4')
                current_project = project
            doc.add_paragraph(f"• {doc_type}: {count} documents", style='List Bullet')
    
    # Save the document
    report_file = "Enhanced_PDF_Extraction_Report.docx"
    doc.save(report_file)
    print(f"Enhanced Word report generated: {os.path.abspath(report_file)}")
    
    return report_file

def main():
    # Connect to database
    connection = connect_to_database()
    if not connection:
        print("Could not connect to database. Exiting.")
        return
    
    # Get enhanced statistics
    print("Gathering enhanced statistics...")
    stats = get_enhanced_stats(connection)
    
    # Close connection
    connection.close()
    
    # Create enhanced charts
    print("Creating enhanced charts...")
    charts = create_enhanced_charts(stats)
    
    # Generate enhanced Word report
    print("Generating enhanced Word report...")
    report_file = generate_enhanced_report(stats, charts)
    
    print("Done!")

if __name__ == "__main__":
    main()
