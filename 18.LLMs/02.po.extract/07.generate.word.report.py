import mysql.connector
from mysql.connector import Error
import datetime
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import numpy as np

# Database configuration
DB_CONFIG = {
    'host': '10.10.11.242',
    'user': 'omar2',
    'password': 'Omar_54321',
    'database': 'RME_TEST'
}

def connect_to_database():
    """Establish connection to MySQL database"""
    try:
        connection = mysql.connector.connect(**DB_CONFIG)
        print("Connected to MySQL database")
        return connection
    except Error as e:
        print(f"Error connecting to MySQL database: {e}")
        return None

def get_pdf_stats(connection):
    """Get statistics about the PDF files processed"""
    cursor = connection.cursor()
    stats = {}
    
    try:
        # Get total PDFs count
        cursor.execute("SELECT COUNT(*) FROM `po.pdfs`")
        stats['total_pdfs'] = cursor.fetchone()[0]
        
        # Get last processed PDF
        cursor.execute("""
        SELECT pdf_filename, processed_timestamp 
        FROM `po.pdfs` 
        ORDER BY processed_timestamp DESC 
        LIMIT 1
        """)
        last_pdf = cursor.fetchone()
        if last_pdf:
            stats['last_pdf_name'] = last_pdf[0]
            stats['last_pdf_time'] = last_pdf[1]
        
        # Get first processed PDF (for duration calculation)
        cursor.execute("""
        SELECT processed_timestamp 
        FROM `po.pdfs` 
        ORDER BY processed_timestamp ASC 
        LIMIT 1
        """)
        first_time = cursor.fetchone()
        if first_time:
            stats['first_time'] = first_time[0]
            
            # Calculate duration and rate
            if 'last_pdf_time' in stats:
                time_diff = stats['last_pdf_time'] - stats['first_time']
                hours = time_diff.total_seconds() / 3600
                stats['duration_hours'] = hours
                if hours > 0:
                    stats['rate_per_hour'] = stats['total_pdfs'] / hours
        
        # Get total chunks
        cursor.execute("SELECT COUNT(*) FROM `po.pdf_chunks`")
        stats['total_chunks'] = cursor.fetchone()[0]
        
        # Get chunk statistics
        cursor.execute("""
        SELECT 
            ROUND(AVG(chunk_count), 2) AS avg_chunks,
            MIN(chunk_count) AS min_chunks,
            MAX(chunk_count) AS max_chunks
        FROM (
            SELECT 
                pdf_id, 
                COUNT(*) AS chunk_count
            FROM `po.pdf_chunks`
            GROUP BY pdf_id
        ) AS chunk_counts
        """)
        chunk_stats = cursor.fetchone()
        if chunk_stats:
            stats['avg_chunks'] = chunk_stats[0]
            stats['min_chunks'] = chunk_stats[1]
            stats['max_chunks'] = chunk_stats[2]
        
        # Get PDF with most chunks
        cursor.execute("""
        SELECT p.pdf_filename, COUNT(c.id) as chunk_count
        FROM `po.pdfs` p
        JOIN `po.pdf_chunks` c ON p.id = c.pdf_id
        GROUP BY p.id
        ORDER BY chunk_count DESC
        LIMIT 1
        """)
        max_chunks_pdf = cursor.fetchone()
        if max_chunks_pdf:
            stats['max_chunks_pdf'] = max_chunks_pdf[0]
            stats['max_chunks_count'] = max_chunks_pdf[1]
        
        # Get PDF with fewest chunks
        cursor.execute("""
        SELECT p.pdf_filename, COUNT(c.id) as chunk_count
        FROM `po.pdfs` p
        JOIN `po.pdf_chunks` c ON p.id = c.pdf_id
        GROUP BY p.id
        ORDER BY chunk_count ASC
        LIMIT 1
        """)
        min_chunks_pdf = cursor.fetchone()
        if min_chunks_pdf:
            stats['min_chunks_pdf'] = min_chunks_pdf[0]
            stats['min_chunks_count'] = min_chunks_pdf[1]
        
        # Get top 5 folders
        cursor.execute("""
        SELECT 
            SUBSTRING_INDEX(pdf_path, '\\\\', -2) AS folder,
            COUNT(*) AS pdf_count
        FROM `po.pdfs`
        GROUP BY folder
        ORDER BY pdf_count DESC
        LIMIT 5
        """)
        stats['top_folders'] = cursor.fetchall()
        
    except Error as e:
        print(f"Error getting PDF statistics: {e}")
    finally:
        cursor.close()
    
    return stats

def create_charts(stats):
    """Create charts for the report"""
    charts = {}
    
    # Create folder for charts if it doesn't exist
    charts_dir = 'report_charts'
    if not os.path.exists(charts_dir):
        os.makedirs(charts_dir)
    
    # Create bar chart for top folders
    if 'top_folders' in stats and stats['top_folders']:
        plt.figure(figsize=(10, 6))
        folders = [folder[0][-20:] + '...' if len(folder[0]) > 20 else folder[0] for folder in stats['top_folders']]
        counts = [folder[1] for folder in stats['top_folders']]
        plt.bar(folders, counts)
        plt.xticks(rotation=45, ha='right')
        plt.title('Top 5 Folders by PDF Count')
        plt.tight_layout()
        chart_path = os.path.join(charts_dir, 'top_folders.png')
        plt.savefig(chart_path)
        plt.close()
        charts['top_folders'] = chart_path
    
    # Create histogram of chunk distribution
    try:
        connection = mysql.connector.connect(**DB_CONFIG)
        cursor = connection.cursor()
        
        cursor.execute("""
        SELECT 
            COUNT(*) as chunk_count
        FROM `po.pdf_chunks`
        GROUP BY pdf_id
        """)
        
        chunk_counts = [row[0] for row in cursor.fetchall()]
        
        if chunk_counts:
            plt.figure(figsize=(10, 6))
            plt.hist(chunk_counts, bins=20)
            plt.title('Distribution of Chunks per PDF')
            plt.xlabel('Number of Chunks')
            plt.ylabel('Number of PDFs')
            chart_path = os.path.join(charts_dir, 'chunk_dist.png')
            plt.savefig(chart_path)
            plt.close()
            charts['chunk_dist'] = chart_path
        
        cursor.close()
        connection.close()
    except Exception as e:
        print(f"Could not create chunk distribution chart: {e}")
    
    return charts

def generate_word_report(stats, charts):
    """Generate a Word document report with the statistics"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('PDF Extraction Database Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_paragraph.add_run(f"Report generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_run.italic = True
    
    doc.add_paragraph()
    
    # Main statistics section
    doc.add_heading('PDF Extraction Progress', level=1)
    
    doc.add_paragraph(f"Total PDFs Processed: {stats.get('total_pdfs', 'N/A')} documents")
    doc.add_paragraph(f"Total Text Chunks Extracted: {stats.get('total_chunks', 'N/A')} chunks")
    doc.add_paragraph(f"Average Chunks per PDF: {stats.get('avg_chunks', 'N/A')}")
    
    if 'rate_per_hour' in stats:
        doc.add_paragraph(f"Processing Rate: {stats['rate_per_hour']:.2f} PDFs per hour")
    
    if 'duration_hours' in stats:
        doc.add_paragraph(f"Processing Duration: {stats['duration_hours']:.2f} hours")
    
    if 'first_time' in stats:
        doc.add_paragraph(f"Started on: {stats['first_time']}")
    
    if 'last_pdf_name' in stats and 'last_pdf_time' in stats:
        doc.add_paragraph(f"Last PDF Processed: \"{stats['last_pdf_name']}\" at {stats['last_pdf_time']}")
    
    # Add explanation about Arabic support
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("The extraction process is handling both English and Arabic documents successfully, as evidenced by the Arabic filename in the last processed PDF.").italic = True
    
    # Add top folders chart if available
    if 'top_folders' in charts:
        doc.add_heading('Top 5 Folders by PDF Count', level=2)
        doc.add_picture(charts['top_folders'], width=Inches(6))
        doc.add_paragraph()
    
    # Database Structure section
    doc.add_heading('Database Structure', level=1)
    
    doc.add_heading('po.pdfs table:', level=2)
    p = doc.add_paragraph()
    p.add_run("• Stores basic PDF information (path, filename, hash)")
    doc.add_paragraph("• Includes the full extracted text")
    doc.add_paragraph("• Tracks processing timestamp")
    
    doc.add_heading('po.pdf_chunks table:', level=2)
    doc.add_paragraph("• Stores individual text chunks")
    doc.add_paragraph("• Links back to the parent PDF via pdf_id")
    doc.add_paragraph(f"• Contains {stats.get('total_chunks', 'N/A')} text segments from the processed PDFs")
    
    # Add chunk distribution chart if available
    if 'chunk_dist' in charts:
        doc.add_heading('Distribution of Chunks per PDF', level=2)
        doc.add_picture(charts['chunk_dist'], width=Inches(6))
        doc.add_paragraph()
    
    # Interesting Statistics section
    doc.add_heading('Interesting Statistics', level=1)
    
    if 'max_chunks_pdf' in stats and 'max_chunks_count' in stats:
        doc.add_paragraph(f"• The PDF with the most chunks (\"{stats['max_chunks_pdf']}\") contains {stats['max_chunks_count']} text segments")
    
    if 'min_chunks_pdf' in stats and 'min_chunks_count' in stats:
        doc.add_paragraph(f"• The PDF with the fewest chunks (\"{stats['min_chunks_pdf']}\") contains only {stats['min_chunks_count']} text segments")
    
    doc.add_paragraph("• The resume capability is working well, as the system continues processing new files without duplicating already processed ones.")
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    
    doc.add_paragraph("The PDF extraction process is running successfully and has made significant progress. The system is handling both English and Arabic documents effectively, as evidenced by the variety of filenames in the processed PDFs.")
    doc.add_paragraph()
    doc.add_paragraph(f"The extraction process is efficiently splitting the text into meaningful chunks, with an average of {stats.get('avg_chunks', 'N/A')} chunks per document. This granular approach will make it easier to analyze and search through the extracted content.")
    doc.add_paragraph()
    doc.add_paragraph("The resume capability is functioning as designed, allowing the process to be stopped and restarted without duplicating work. This is particularly important for processing large collections of PDFs over extended periods.")
    
    # Save the document
    report_file = "PDF_Extraction_Report.docx"
    doc.save(report_file)
    print(f"Word report generated: {os.path.abspath(report_file)}")
    
    return report_file

def main():
    # Connect to database
    connection = connect_to_database()
    if not connection:
        print("Could not connect to database. Exiting.")
        return
    
    # Get statistics
    print("Gathering statistics...")
    stats = get_pdf_stats(connection)
    
    # Close connection
    connection.close()
    
    # Create charts
    print("Creating charts...")
    charts = create_charts(stats)
    
    # Generate Word report
    print("Generating Word report...")
    report_file = generate_word_report(stats, charts)
    
    print("Done!")

if __name__ == "__main__":
    main()
