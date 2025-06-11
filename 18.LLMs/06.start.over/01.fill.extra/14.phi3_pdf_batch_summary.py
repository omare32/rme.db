import os
import subprocess
from PyPDF2 import PdfReader
from docx import Document
import time

PDF_SOURCE_FOLDER = r'H:\Projects Control (PC)\10 Backup\06 Yasser\Damietta Buildings Project'
OUTPUT_DIR = r'D:\OEssam\Test\phi3'
PROCESSED_FILES_LOG = os.path.join(OUTPUT_DIR, 'processed_pdfs.txt')
LLM_MODEL_TAG = 'phi3:latest'

SUMMARY_PROMPT = """Summarize the following document. Focus only on the most essential information, such as project name, contract value, parties, dates, and any key terms or obligations. Be concise and clear.\n\nDocument Text:\n{text}\n\nSummary:"""

os.makedirs(OUTPUT_DIR, exist_ok=True)

def get_processed_pdfs():
    if not os.path.exists(PROCESSED_FILES_LOG):
        return set()
    with open(PROCESSED_FILES_LOG, 'r', encoding='utf-8') as f:
        return set(line.strip() for line in f)

def mark_pdf_as_processed(pdf_path):
    with open(PROCESSED_FILES_LOG, 'a', encoding='utf-8') as f:
        f.write(pdf_path + '\n')

def get_all_unprocessed_pdfs():
    processed = get_processed_pdfs()
    pdfs = []
    for dirpath, _, filenames in os.walk(PDF_SOURCE_FOLDER):
        for fn in filenames:
            if fn.lower().endswith('.pdf'):
                full = os.path.join(dirpath, fn)
                if full not in processed:
                    pdfs.append(full)
    pdfs.sort()
    return pdfs

def extract_text_from_pdf(pdf_path):
    try:
        reader = PdfReader(pdf_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")
        return "\n".join(text)
    except Exception as e:
        print(f"[ERROR] Failed to extract text from {pdf_path}: {e}")
        return ""

def query_phi3(prompt):
    start = time.time()
    process = subprocess.Popen([
        "ollama", "run", LLM_MODEL_TAG
    ], stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    stdout, stderr = process.communicate(prompt.encode("utf-8"))
    duration = time.time() - start
    out = stdout.decode("utf-8", errors="ignore")
    err = stderr.decode("utf-8", errors="ignore")
    if err:
        print(f"[phi3 ERROR]: {err}")
    return out.strip(), duration

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    if not os.path.exists(PROCESSED_FILES_LOG):
        with open(PROCESSED_FILES_LOG, 'w', encoding='utf-8') as f:
            pass
    pdfs = get_all_unprocessed_pdfs()
    if not pdfs:
        print("No new PDFs to process.")
        return
    for pdf_path in pdfs:
        print(f"\n=== Processing: {pdf_path} ===")
        text = extract_text_from_pdf(pdf_path)
        if not text.strip():
            print(f"[WARNING] No extractable text in {pdf_path}. Skipping.")
            mark_pdf_as_processed(pdf_path)
            continue
        prompt = SUMMARY_PROMPT.format(text=text[:8000])  # Limit to 8k chars for speed/safety
        summary, duration = query_phi3(prompt)
        print(f"Summary generated in {duration:.1f}s.")
        doc = Document()
        doc.add_heading(f'Phi3 Summary: {os.path.basename(pdf_path)}', 0)
        doc.add_paragraph(f'Source PDF: {pdf_path}')
        doc.add_paragraph(summary)
        out_name = f"phi3_summary_{os.path.splitext(os.path.basename(pdf_path))[0]}.docx"
        out_path = os.path.join(OUTPUT_DIR, out_name)
        try:
            doc.save(out_path)
            print(f"Saved summary to {out_path}")
            mark_pdf_as_processed(pdf_path)
        except Exception as e:
            print(f"[ERROR] Could not save Word file for {pdf_path}: {e}")

if __name__ == "__main__":
    main()
