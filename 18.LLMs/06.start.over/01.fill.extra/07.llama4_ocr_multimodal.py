import os
import base64
import subprocess
from PIL import Image
from pdf2image import convert_from_path
import pytesseract
import cv2
import numpy as np
from scipy.ndimage import interpolation
from docx import Document

# === CONFIGURATION ===
import pymysql
# MySQL connection settings (reuse from previous script)
HOST = '10.10.11.242'
USER = 'omar2'
PASSWORD = 'Omar_54321'
DB = 'RME_TEST'
TABLE = 'po.pdfs'

POPPLER_Path = r"C:\\Program Files\\poppler\\Library\\bin"
pytesseract.pytesseract.tesseract_cmd = r"C:\\Users\\Omar Essam2\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe"
LLM_MODEL = "llama4:latest"  # Use the model tag available in your Ollama
PROMPT_TEMPLATE = """You are an OCR agent. Below is the raw text extracted using Tesseract OCR. Use it along with the image to accurately extract relevant Purchase Order (PO), total amount, date and most importantly terms and conditions (Arabic and English if present).\n\nRaw OCR Text:\n{text}\n"""

def get_random_pdf_paths(n=3):
    conn = pymysql.connect(host=HOST, user=USER, password=PASSWORD, database=DB, charset='utf8mb4', cursorclass=pymysql.cursors.DictCursor)
    with conn.cursor() as cursor:
        cursor.execute(f"SELECT pdf_path FROM `{TABLE}` WHERE pdf_path IS NOT NULL AND pdf_path != '' AND document_type = 'Purchase Order'")
        rows = cursor.fetchall()
    conn.close()
    import random
    all_paths = [row['pdf_path'] for row in rows]
    if not all_paths:
        return []
    return random.sample(all_paths, min(n, len(all_paths)))

# === Step 1: Convert PDF to images ===
def pdf_to_images(pdf_path, output_dir, max_pages=3):
    images = convert_from_path(pdf_path, dpi=300, poppler_path=POPPLER_Path, first_page=1, last_page=max_pages)
    image_paths = []
    for i, img in enumerate(images):
        out_name = f"{os.path.splitext(os.path.basename(pdf_path))[0]}_page{i+1}.png"
        path = os.path.join(output_dir, out_name)
        img.save(path, "PNG")
        image_paths.append(path)
    return image_paths

# === Step 2: Preprocess image for OCR ===
def preprocess_image(image_path):
    image = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
    # 1. Increase contrast using histogram equalization
    image = cv2.equalizeHist(image)
    # 2. Adaptive thresholding
    image = cv2.adaptiveThreshold(image, 255,
                                  cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                  cv2.THRESH_BINARY, 35, 11)
    # 3. Remove noise
    image = cv2.medianBlur(image, 3)
    # 4. Deskew (rotate to correct tilt)
    coords = np.column_stack(np.where(image > 0))
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    (h, w) = image.shape
    M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
    image = cv2.warpAffine(image, M, (w, h),
                           flags=cv2.INTER_CUBIC,
                           borderMode=cv2.BORDER_REPLICATE)
    # Save processed image (overwrite or use temp)
    processed_path = image_path.replace(".png", "_processed.png")
    cv2.imwrite(processed_path, image)
    return processed_path

# === Step 3: Run Tesseract OCR ===
def run_tesseract(image_path):
    try:
        return pytesseract.image_to_string(Image.open(image_path), lang='ara+eng')
    except Exception as e:
        print(f"Tesseract error: {e}")
        return ""

# === Step 4: Convert image to base64 ===
def image_to_base64(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode("utf-8")

# === Step 5: Query LLaMA 4 (Ollama) ===
def query_llama4(image_b64, raw_text):
    prompt = PROMPT_TEMPLATE.format(text=raw_text)
    full_prompt = f"<|image|>{image_b64}\n{prompt}"
    process = subprocess.Popen(
        ["ollama", "run", LLM_MODEL],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    stdout, stderr = process.communicate(full_prompt.encode("utf-8"))
    stdout = stdout.decode("utf-8", errors="ignore")
    stderr = stderr.decode("utf-8", errors="ignore")
    if stderr:
        print("Error:", stderr)
    return stdout

# === Main Pipeline ===
def main():
    output_dir = r'D:\OEssam\Test\llama4'
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    doc.add_heading('LLaMA4 Multimodal OCR Results', 0)
    pdf_paths = get_random_pdf_paths(3)
    if not pdf_paths:
        print("No PDF paths found in database!")
        return
    for pdf_idx, PDF_PATH in enumerate(pdf_paths, 1):
        doc.add_heading(f'PDF {pdf_idx}: {PDF_PATH}', level=1)
        print(f"\n=== Processing PDF {pdf_idx}: {PDF_PATH} ===")
        if not os.path.exists(PDF_PATH):
            print(f"[ERROR] PDF not found: {PDF_PATH}")
            doc.add_paragraph("[ERROR] PDF not found!")
            continue
        image_files = pdf_to_images(PDF_PATH, output_dir, max_pages=3)
        for idx, image_file in enumerate(image_files):
            doc.add_heading(f'Page {idx+1}', level=2)
            print(f"\n--- Processing Page {idx+1} ---")
            processed_image = preprocess_image(image_file)
            raw_ocr = run_tesseract(processed_image)
            print("Raw OCR Preview:", raw_ocr[:300], "...")
            img_b64 = image_to_base64(processed_image)
            response = query_llama4(img_b64, raw_ocr)
            print("LLaMA4 Response:\n", response)
            print("-" * 60)
            # Save to docx
            doc.add_paragraph('Raw OCR Preview:')
            doc.add_paragraph(raw_ocr)
            doc.add_paragraph('LLaMA4 Response:')
            doc.add_paragraph(response)
            doc.add_paragraph('-' * 40)
    result_docx = os.path.join(output_dir, 'llama4_ocr_multimodal_results.docx')
    doc.save(result_docx)
    print(f"\nAll results saved to {result_docx}")

if __name__ == "__main__":
    main()
