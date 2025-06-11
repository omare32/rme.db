import os
import base64
import pytesseract
from PIL import Image
from docx import Document
import subprocess
import time

IMAGE_ROOT = r'D:\OEssam\Test\epico-pdf-images'
OUTPUT_DIR = r'D:\OEssam\Test\gemma3-epico3'
PROCESSED_FILES_LOG = os.path.join(OUTPUT_DIR, 'processed_folders.txt')
LLM_MODEL_TAG = 'gemma3:latest'
pytesseract.pytesseract.tesseract_cmd = r"C:\\Users\\Omar Essam2\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe"

PAGE_PROMPT = """You are an expert at reading scanned documents. Below is the OCR text and the image of a single page. Extract and summarize the essential information (project, contract value, parties, dates, key terms, obligations, etc). Be concise and clear.\n\nOCR Text:\n{text}\n"""

SUMMARY_PROMPT = """You are an expert summarizer. Here are the summaries of all pages in a scanned contract. Write a concise, comprehensive summary of the entire document, focusing on the most important information.\n\nPage Summaries:\n{text}\n\nFull Document Summary:"""

def get_processed_folders():
    if not os.path.exists(PROCESSED_FILES_LOG):
        return set()
    with open(PROCESSED_FILES_LOG, 'r', encoding='utf-8') as f:
        return set(line.strip() for line in f)

def mark_folder_as_processed(folder_name):
    with open(PROCESSED_FILES_LOG, 'a', encoding='utf-8') as f:
        f.write(folder_name + '\n')

def get_all_image_folders():
    return [os.path.join(IMAGE_ROOT, d) for d in os.listdir(IMAGE_ROOT)
            if os.path.isdir(os.path.join(IMAGE_ROOT, d))]

def run_tesseract(image_path):
    try:
        return pytesseract.image_to_string(Image.open(image_path), lang='ara+eng')
    except Exception as e:
        print(f"Tesseract error: {e}")
        return ""

def image_to_base64(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode("utf-8")

def query_gemma3(prompt, image_b64=None):
    if image_b64:
        prompt = f"<|image|>{image_b64}\n{prompt}"
    start = time.time()
    process = subprocess.Popen([
        "ollama", "run", LLM_MODEL_TAG
    ], stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    stdout, stderr = process.communicate(prompt.encode("utf-8"))
    duration = time.time() - start
    out = stdout.decode("utf-8", errors="ignore")
    err = stderr.decode("utf-8", errors="ignore")
    if err:
        print(f"[gemma3 ERROR]: {err}")
    return out.strip(), duration

def process_folder(folder_path):
    doc = Document()
    folder_name = os.path.basename(folder_path)
    doc.add_heading(f'Gemma3 Multimodal Summary: {folder_name}', 0)
    doc.add_paragraph(f'Source Image Folder: {folder_path}')
    print(f"\n=== Processing Folder: {folder_path} ===")
    image_files = sorted([os.path.join(folder_path, f) for f in os.listdir(folder_path)
                         if f.lower().endswith('.png') or f.lower().endswith('.jpg')])
    all_page_summaries = []
    for idx, image_file_path in enumerate(image_files):
        page_num = idx + 1
        doc.add_heading(f'Page {page_num}', level=2)
        print(f"\n--- Processing Page {page_num} of {len(image_files)} ({os.path.basename(image_file_path)}) ---")
        ocr_text = run_tesseract(image_file_path)
        img_b64 = image_to_base64(image_file_path)
        page_prompt = PAGE_PROMPT.format(text=ocr_text)
        page_summary, duration = query_gemma3(page_prompt, img_b64)
        all_page_summaries.append(page_summary)
        doc.add_paragraph(f'Duration: {duration:.2f} seconds')
        doc.add_paragraph('Raw OCR Preview:')
        doc.add_paragraph(ocr_text)
        doc.add_paragraph('Gemma3 Page Summary:')
        doc.add_paragraph(page_summary)
        doc.add_paragraph('-' * 40)
    if all_page_summaries:
        print("\n--- Generating Full Document Summary --- ")
        concatenated = "\n\n".join(all_page_summaries)
        summary_prompt = SUMMARY_PROMPT.format(text=concatenated)
        final_summary, summary_duration = query_gemma3(summary_prompt)
        doc.add_heading('Final Comprehensive Summary', level=1)
        doc.add_paragraph(f'Duration for summary generation: {summary_duration:.2f} seconds')
        doc.add_paragraph(final_summary)
    output_word_filename = f"gemma3_multimodal_summary_{folder_name}.docx"
    result_docx_path = os.path.join(OUTPUT_DIR, output_word_filename)
    try:
        doc.save(result_docx_path)
        print(f"\nResults for {folder_name} saved to {result_docx_path}")
        mark_folder_as_processed(folder_name)
    except Exception as e:
        print(f"[ERROR] Failed to save Word document {result_docx_path}: {e}")

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    if not os.path.exists(PROCESSED_FILES_LOG):
        with open(PROCESSED_FILES_LOG, 'w', encoding='utf-8') as f:
            pass
    folders = get_all_image_folders()
    processed = get_processed_folders()
    to_process = [f for f in folders if os.path.basename(f) not in processed]
    if not to_process:
        print("No new folders to process.")
        return
    for folder_path in to_process:
        process_folder(folder_path)
        print("\n" + "="*50 + "\n")

if __name__ == "__main__":
    main()
