# This script requires mysql-connector-python and python-docx
# Install them using:
# pip install mysql-connector-python python-docx

import os
import mysql.connector
from docx import Document

# === DATABASE CONFIG ===
DB_CONFIG = {
    'host': '10.10.11.242',
    'user': 'omar2',
    'password': 'Omar_54321',
    'database': 'RME_TEST'
}
TABLE_NAME = 'project_summaries'

# === PROJECT FOLDERS ===
PROJECT_FOLDERS = {
    'Damietta Buildings Project': r'D:\OEssam\Test\gemma3-3rd-time',
    'Eipico 3': r'D:\OEssam\Test\gemma3-epico3',
    'R5': r'D:\OEssam\Test\gemma3-r5'
}

# === DATABASE FUNCTIONS ===
def get_db_connection():
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        return conn
    except mysql.connector.Error as err:
        print(f"Error connecting to database: {err}")
        return None

def create_table_if_not_exists(cursor):
    """Creates the table if it doesn't exist, and ensures summary columns are MEDIUMTEXT."""
    # Step 1: Create the table with MEDIUMTEXT if it doesn't exist.
    create_table_query = f"""
    CREATE TABLE IF NOT EXISTS {TABLE_NAME} (
        id INT AUTO_INCREMENT PRIMARY KEY,
        project_name VARCHAR(255) NOT NULL,
        pdf_filename VARCHAR(255) NOT NULL,
        final_summary MEDIUMTEXT,
        page_summaries MEDIUMTEXT,
        UNIQUE KEY unique_project_file (project_name, pdf_filename)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4; 
    """
    cursor.execute(create_table_query)
    
    # Step 2: Alter the table to ensure the columns are MEDIUMTEXT, in case the table
    # already existed with the old TEXT type.
    try:
        alter_query = f"ALTER TABLE {TABLE_NAME} MODIFY COLUMN final_summary MEDIUMTEXT, MODIFY COLUMN page_summaries MEDIUMTEXT;"
        cursor.execute(alter_query)
    except mysql.connector.Error:
        # This error is expected if the columns are already MEDIUMTEXT. We can ignore it.
        pass

def is_record_present(cursor, project_name, pdf_filename):
    query = f"SELECT id FROM {TABLE_NAME} WHERE project_name = %s AND pdf_filename = %s"
    cursor.execute(query, (project_name, pdf_filename))
    return cursor.fetchone() is not None

def insert_record(cursor, project_name, pdf_filename, final_summary, page_summaries):
    query = f"""
    INSERT INTO {TABLE_NAME} (project_name, pdf_filename, final_summary, page_summaries)
    VALUES (%s, %s, %s, %s)
    """
    cursor.execute(query, (project_name, pdf_filename, final_summary, page_summaries))

# === WORD FILE PARSING ===
def parse_word_document(doc_path):
    try:
        doc = Document(doc_path)
        pdf_filename = "Unknown"
        if 'gemma3_multimodal_summary_' in os.path.basename(doc_path):
            pdf_filename = os.path.basename(doc_path).replace('gemma3_multimodal_summary_', '').replace('.docx', '')

        page_summaries = []
        final_summary = ""
        is_final_summary = False

        for para in doc.paragraphs:
            # Check for final summary heading
            if 'Final Comprehensive Summary' in para.text:
                is_final_summary = True
                continue
            
            # Check for page summary heading
            if para.style.name.startswith('Heading 2'): # Assumes 'Page X' uses Heading 2 style
                 is_final_summary = False # Stop capturing final summary if a new page starts

            if is_final_summary:
                final_summary += para.text + '\n'
            else:
                # Crude way to find page summaries by looking for specific text
                if 'Gemma3 Page Summary:' in para.text:
                    # The actual summary is likely the next paragraph
                    pass # Logic needs to be more robust
                elif 'Raw OCR Preview:' in para.text:
                    pass
        
        # A more robust way to get page summaries by looking at the structure
        all_paras = doc.paragraphs
        for i, para in enumerate(all_paras):
            if 'Gemma3 Page Summary:' in para.text and i + 1 < len(all_paras):
                page_summaries.append(all_paras[i+1].text)

        return {
            'pdf_filename': pdf_filename,
            'final_summary': final_summary.strip(),
            'page_summaries': '\n\n---\n\n'.join(page_summaries).strip()
        }
    except Exception as e:
        print(f"Error parsing {doc_path}: {e}")
        return None

# === MAIN EXECUTION ===
def main():
    conn = get_db_connection()
    if not conn:
        return
    
    cursor = conn.cursor()
    create_table_if_not_exists(cursor)
    conn.commit()  # Commit schema changes before proceeding
    print(f"Table '{TABLE_NAME}' is ready.")

    total_inserted = 0
    for project_name, folder_path in PROJECT_FOLDERS.items():
        print(f"\nProcessing project: {project_name}")
        if not os.path.isdir(folder_path):
            print(f"  [WARNING] Folder not found: {folder_path}")
            continue

        for filename in os.listdir(folder_path):
            if filename.endswith('.docx'):
                doc_path = os.path.join(folder_path, filename)
                data = parse_word_document(doc_path)
                
                if data and not is_record_present(cursor, project_name, data['pdf_filename']):
                    print(f"  Inserting data for: {data['pdf_filename']}")
                    insert_record(cursor, project_name, data['pdf_filename'], data['final_summary'], data['page_summaries'])
                    total_inserted += 1
                elif data:
                    print(f"  Skipping existing record: {data['pdf_filename']}")

    conn.commit()
    cursor.close()
    conn.close()
    print(f"\nDatabase operation complete. Inserted {total_inserted} new records.")

if __name__ == "__main__":
    main()
