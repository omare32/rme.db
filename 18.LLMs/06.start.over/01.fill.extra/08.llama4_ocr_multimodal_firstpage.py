import os
import base64
import subprocess
from PIL import Image
from pdf2image import convert_from_path
import pytesseract
import cv2
import numpy as np
from scipy.ndimage import interpolation
from docx import Document
import pymysql

# MySQL connection settings
HOST = '10.10.11.242'
USER = 'omar2'
PASSWORD = 'Omar_54321'
DB = 'RME_TEST'
TABLE = 'po.pdfs'

POPPLER_Path = r"C:\\Program Files\\poppler\\Library\\bin"
pytesseract.pytesseract.tesseract_cmd = r"C:\\Users\\Omar Essam2\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe"
LLM_MODELS = [
    ("llama4:latest", "Llama-4"),
    ("qwen2.5:32b", "Qwen2.5-32B"),
    ("gemma3:latest", "Gemma3"),
    ("qwen2.5vl:latest", "Qwen2.5VL"),
    ("phi3:latest", "Phi3")
]
PROMPT_TEMPLATE = """You are an OCR agent. Below is the raw text extracted using Tesseract OCR. Use it along with the image to accurately extract relevant Purchase Order (PO), total amount, date and most importantly terms and conditions (Arabic and English if present).\n\nRaw OCR Text:\n{text}\n"""

def get_random_pdf_paths(n=3):
    conn = pymysql.connect(host=HOST, user=USER, password=PASSWORD, database=DB, charset='utf8mb4', cursorclass=pymysql.cursors.DictCursor)
    with conn.cursor() as cursor:
        cursor.execute(f"SELECT pdf_path FROM `{TABLE}` WHERE pdf_path IS NOT NULL AND pdf_path != '' AND document_type = 'Purchase Order'")
        rows = cursor.fetchall()
    conn.close()
    import random
    all_paths = [row['pdf_path'] for row in rows]
    if not all_paths:
        return []
    return random.sample(all_paths, min(n, len(all_paths)))

def pdf_first_page(pdf_path, output_dir):
    images = convert_from_path(pdf_path, dpi=300, poppler_path=POPPLER_Path, first_page=1, last_page=1)
    if images:
        out_name = f"{os.path.splitext(os.path.basename(pdf_path))[0]}_page1.png"
        path = os.path.join(output_dir, out_name)
        images[0].save(path, "PNG")
        return path
    return None

def preprocess_image(image_path):
    image = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
    image = cv2.equalizeHist(image)
    image = cv2.adaptiveThreshold(image, 255,
                                  cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                  cv2.THRESH_BINARY, 35, 11)
    image = cv2.medianBlur(image, 3)
    coords = np.column_stack(np.where(image > 0))
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    (h, w) = image.shape
    M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
    image = cv2.warpAffine(image, M, (w, h),
                           flags=cv2.INTER_CUBIC,
                           borderMode=cv2.BORDER_REPLICATE)
    processed_path = image_path.replace(".png", "_processed.png")
    cv2.imwrite(processed_path, image)
    return processed_path

def run_tesseract(image_path):
    try:
        return pytesseract.image_to_string(Image.open(image_path), lang='ara+eng')
    except Exception as e:
        print(f"Tesseract error: {e}")
        return ""

def image_to_base64(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode("utf-8")

import time

def query_model(image_b64, raw_text, model_tag):
    prompt = PROMPT_TEMPLATE.format(text=raw_text)
    full_prompt = f"<|image|>{image_b64}\n{prompt}"
    start_time = time.time()
    process = subprocess.Popen(
        ["ollama", "run", model_tag],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    stdout, stderr = process.communicate(full_prompt.encode("utf-8"))
    duration = time.time() - start_time
    stdout = stdout.decode("utf-8", errors="ignore")
    stderr = stderr.decode("utf-8", errors="ignore")
    if stderr:
        print(f"Error ({model_tag}):", stderr)
    return stdout, duration

def main():
    output_dir = r'D:\OEssam\Test\llama4'
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    doc.add_heading('Multimodal OCR Results (First Page, 3 Models)', 0)

    pdf_paths = get_random_pdf_paths(3)
    if not pdf_paths:
        print("No PDF paths found in database!")
        return
    # For duration summary
    durations = []
    results = []
    for pdf_idx, PDF_PATH in enumerate(pdf_paths, 1):
        doc.add_heading(f'PDF {pdf_idx}: {PDF_PATH}', level=1)
        print(f"\n=== Processing PDF {pdf_idx}: {PDF_PATH} ===")
        if not os.path.exists(PDF_PATH):
            print(f"[ERROR] PDF not found: {PDF_PATH}")
            doc.add_paragraph("[ERROR] PDF not found!")
            continue
        image_file = pdf_first_page(PDF_PATH, output_dir)
        if not image_file:
            print(f"[ERROR] Could not convert first page: {PDF_PATH}")
            doc.add_paragraph("[ERROR] Could not convert first page!")
            continue
        doc.add_heading('Page 1', level=2)
        print(f"--- Processing First Page ---")
        processed_image = preprocess_image(image_file)
        raw_ocr = run_tesseract(processed_image)
        print("Raw OCR Preview:", raw_ocr[:300], "...")
        img_b64 = image_to_base64(processed_image)
        for model_tag, model_name in LLM_MODELS:
            print(f"Running {model_name} ({model_tag})...")
            response, duration = query_model(img_b64, raw_ocr, model_tag)
            print(f"{model_name} Response:\n", response)
            print(f"Duration: {duration:.2f} seconds\n")
            durations.append({
                'pdf_idx': pdf_idx,
                'pdf_path': PDF_PATH,
                'model': model_name,
                'duration': duration
            })
            results.append({
                'pdf_idx': pdf_idx,
                'pdf_path': PDF_PATH,
                'model': model_name,
                'raw_ocr': raw_ocr,
                'response': response
            })
            doc.add_heading(f'{model_name} Result', level=3)
            doc.add_paragraph(f'Duration: {duration:.2f} seconds')
            doc.add_paragraph('Raw OCR Preview:')
            doc.add_paragraph(raw_ocr)
            doc.add_paragraph(f'{model_name} Response:')
            doc.add_paragraph(response)
            doc.add_paragraph('-' * 40)
    # Add summary table at top
    doc.add_page_break()
    doc._body._element.insert(1, doc._body._element[-1])  # Move last page (summary) to top
    doc.paragraphs.insert(1, doc.add_heading('Duration Summary (seconds)', level=1))
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'PDF #'
    hdr_cells[1].text = 'PDF Path'
    hdr_cells[2].text = 'Model'
    hdr_cells[3].text = 'Duration (s)'
    for d in durations:
        row_cells = table.add_row().cells
        row_cells[0].text = str(d['pdf_idx'])
        row_cells[1].text = os.path.basename(d['pdf_path'])
        row_cells[2].text = d['model']
        row_cells[3].text = f"{d['duration']:.2f}"
    result_docx = os.path.join(output_dir, 'llama4_ocr_multimodal_results_firstpage.docx')
    doc.save(result_docx)
    print(f"\nAll results saved to {result_docx}")

if __name__ == "__main__":
    main()
