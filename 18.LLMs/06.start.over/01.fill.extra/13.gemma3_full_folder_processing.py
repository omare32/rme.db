import os
import base64
import subprocess
from PIL import Image
from pdf2image import convert_from_path
import pytesseract
import cv2
import numpy as np
from scipy.ndimage import interpolation
from docx import Document
import time

# === CONFIGURATION ===
POPPLER_Path = r"C:\\Program Files\\poppler\\Library\\bin"
pytesseract.pytesseract.tesseract_cmd = r"C:\\Users\\Omar Essam2\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe"

LLM_MODEL_TAG = "gemma3:latest"
LLM_MODEL_NAME = "Gemma3"

PDF_SOURCE_FOLDER = r'H:\Projects Control (PC)\10 Backup\06 Yasser\Damietta Buildings Project'
OUTPUT_DIR = r'D:\OEssam\Test\gemma3'
IMAGE_OUTPUT_DIR = os.path.join(OUTPUT_DIR, 'images')
PROCESSED_FILES_LOG = os.path.join(OUTPUT_DIR, 'processed_pdfs.txt')

PROMPT_TEMPLATE_PAGE = """You are an OCR agent. Below is the raw text extracted using Tesseract OCR. Use it along with the image to accurately extract relevant Purchase Order (PO) details, total amount, date, and most importantly terms and conditions (Arabic and English if present) from THIS PAGE ONLY.\n\nRaw OCR Text from this page:\n{text}\n\nExtracted Information from this page:"""

PROMPT_TEMPLATE_SUMMARY = """You are an expert document summarizer. The document was found at the following path, which may contain useful context about the supplier or project: {pdf_path}\n\nBased on the following extracted text from all pages of a Purchase Order, please provide a comprehensive summary of all important information. This includes, but is not limited to: Purchase Order (PO) number, total amount, date, supplier details, buyer details, item descriptions, quantities, unit prices, total prices, and all terms and conditions (Arabic and English if present). Consolidate information and present it clearly.\n\nExtracted Text from All Pages:\n{text}\n\nFinal Comprehensive Summary:"""

def get_processed_pdfs():
    if not os.path.exists(PROCESSED_FILES_LOG):
        return set()
    with open(PROCESSED_FILES_LOG, 'r', encoding='utf-8') as f:
        return set(line.strip() for line in f)

def mark_pdf_as_processed(pdf_path):
    with open(PROCESSED_FILES_LOG, 'a', encoding='utf-8') as f:
        f.write(pdf_path + '\n')

def get_all_unprocessed_pdfs():
    if not os.path.isdir(PDF_SOURCE_FOLDER):
        return []
    processed_pdfs = get_processed_pdfs()
    all_pdfs_in_folder = []
    for dirpath, _, filenames in os.walk(PDF_SOURCE_FOLDER):
        for filename in filenames:
            if filename.lower().endswith('.pdf'):
                full_path = os.path.join(dirpath, filename)
                if full_path not in processed_pdfs:
                    all_pdfs_in_folder.append(full_path)
    all_pdfs_in_folder.sort()
    return all_pdfs_in_folder

def pdf_to_images(pdf_path, image_save_dir):
    os.makedirs(image_save_dir, exist_ok=True)
    images = convert_from_path(pdf_path, dpi=300, poppler_path=POPPLER_Path)
    image_paths = []
    pdf_basename = os.path.splitext(os.path.basename(pdf_path))[0]
    for i, img in enumerate(images):
        out_name = f"{pdf_basename}_page{i+1}.png"
        path = os.path.join(image_save_dir, out_name)
        img.save(path, "PNG")
        image_paths.append(path)
    return image_paths

def preprocess_image(image_path):
    image = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
    image = cv2.equalizeHist(image)
    image = cv2.adaptiveThreshold(image, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 35, 11)
    image = cv2.medianBlur(image, 3)
    coords = np.column_stack(np.where(image > 0))
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    (h, w) = image.shape[:2]
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, angle, 1.0)
    image = cv2.warpAffine(image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
    processed_path = os.path.join(os.path.dirname(image_path), f"{os.path.splitext(os.path.basename(image_path))[0]}_processed.png")
    cv2.imwrite(processed_path, image)
    return processed_path

def run_tesseract(image_path):
    try:
        return pytesseract.image_to_string(Image.open(image_path), lang='ara+eng')
    except Exception as e:
        print(f"Tesseract error: {e}")
        return ""

def image_to_base64(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode("utf-8")

def query_gemma3(full_prompt, image_b64=None):
    if image_b64:
        full_prompt = f"<|image|>{image_b64}\n{full_prompt}"
    start_time = time.time()
    process = subprocess.Popen(
        ["ollama", "run", LLM_MODEL_TAG],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    stdout, stderr = process.communicate(full_prompt.encode("utf-8"))
    duration = time.time() - start_time
    stdout = stdout.decode("utf-8", errors="ignore")
    stderr = stderr.decode("utf-8", errors="ignore")
    if stderr:
        print(f"Error ({LLM_MODEL_TAG}):", stderr)
    return stdout, duration

def process_pdf(PDF_PATH):
    doc = Document()
    pdf_filename_base = os.path.splitext(os.path.basename(PDF_PATH))[0]
    doc.add_heading(f'{LLM_MODEL_NAME} Full PDF Processing: {pdf_filename_base}', 0)
    doc.add_paragraph(f'Source PDF: {PDF_PATH}')
    print(f"\n=== Processing PDF: {PDF_PATH} ===")
    if not os.path.exists(PDF_PATH):
        print(f"[ERROR] PDF not found or inaccessible: {PDF_PATH}")
        return
    image_files = pdf_to_images(PDF_PATH, IMAGE_OUTPUT_DIR)
    all_page_gemma_responses = []
    for idx, image_file_path in enumerate(image_files):
        page_num = idx + 1
        doc.add_heading(f'Page {page_num}', level=2)
        print(f"\n--- Processing Page {page_num} of {len(image_files)} ({os.path.basename(image_file_path)}) ---")
        try:
            processed_image_path = preprocess_image(image_file_path)
            raw_ocr = run_tesseract(processed_image_path)
            print(f"Raw OCR Preview (Page {page_num}):", raw_ocr[:200] + "..." if len(raw_ocr) > 200 else raw_ocr)
            img_b64 = image_to_base64(processed_image_path)
            page_prompt = PROMPT_TEMPLATE_PAGE.format(text=raw_ocr)
            gemma_page_response, duration = query_gemma3(page_prompt, img_b64)
            all_page_gemma_responses.append(gemma_page_response)
            print(f"{LLM_MODEL_NAME} Response (Page {page_num}):\n", gemma_page_response[:300] + "..." if len(gemma_page_response) > 300 else gemma_page_response)
            print(f"Duration for page {page_num}: {duration:.2f} seconds")
            doc.add_paragraph(f'Duration: {duration:.2f} seconds')
            doc.add_paragraph('Raw OCR Preview:')
            doc.add_paragraph(raw_ocr)
            doc.add_paragraph(f'{LLM_MODEL_NAME} Page Response:')
            doc.add_paragraph(gemma_page_response)
        except Exception as e:
            print(f"[ERROR] Processing page {page_num} ({os.path.basename(image_file_path)}): {e}")
            doc.add_paragraph(f"[ERROR] Could not process page {page_num} ({os.path.basename(image_file_path)}): {e}")
        doc.add_paragraph('-' * 40)
    if all_page_gemma_responses:
        print("\n--- Generating Final Summary --- ")
        concatenated_responses = "\n\n".join(all_page_gemma_responses)
        summary_prompt = PROMPT_TEMPLATE_SUMMARY.format(pdf_path=PDF_PATH, text=concatenated_responses)
        final_summary, summary_duration = query_gemma3(summary_prompt)
        print(f"Final {LLM_MODEL_NAME} Summary:\n", final_summary)
        print(f"Duration for final summary: {summary_duration:.2f} seconds")
        doc.add_heading('Final Comprehensive Summary', level=1)
        doc.add_paragraph(f'Duration for summary generation: {summary_duration:.2f} seconds')
        doc.add_paragraph(final_summary)
    else:
        print("No page responses to summarize for this PDF.")
        doc.add_heading('Final Comprehensive Summary', level=1)
        doc.add_paragraph("No page-level extractions were successful for this PDF to generate a summary.")
    output_word_filename = f"gemma3_summary_{pdf_filename_base}.docx"
    result_docx_path = os.path.join(OUTPUT_DIR, output_word_filename)
    try:
        doc.save(result_docx_path)
        print(f"\nResults for {os.path.basename(PDF_PATH)} saved to {result_docx_path}")
        mark_pdf_as_processed(PDF_PATH)
    except Exception as e:
        print(f"[ERROR] Failed to save Word document {result_docx_path}: {e}")

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(IMAGE_OUTPUT_DIR, exist_ok=True)
    if not os.path.exists(PROCESSED_FILES_LOG):
        with open(PROCESSED_FILES_LOG, 'w', encoding='utf-8') as f:
            pass
    unprocessed_pdfs = get_all_unprocessed_pdfs()
    if not unprocessed_pdfs:
        print("All PDFs in the source folder have been processed or folder is inaccessible/empty.")
        return
    for pdf_path in unprocessed_pdfs:
        process_pdf(pdf_path)
        print("\n" + "="*50 + "\n")

if __name__ == "__main__":
    main()
