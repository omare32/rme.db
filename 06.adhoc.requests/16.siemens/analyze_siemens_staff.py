import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from datetime import datetime

# Read the Excel file
df = pd.read_excel('Siemens Staff.xlsx')

# Create a new Word document
doc = Document()
doc.add_heading('Siemens Staff Analysis Report', 0)
doc.add_paragraph(f'Report generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

# 1. Total number of unique positions
unique_positions = df['Job'].nunique()
doc.add_heading('1. Position Overview', level=1)
doc.add_paragraph(f'Total number of unique positions: {unique_positions}')

# 2. Top 10 positions
top_positions = df['Job'].value_counts().head(10)
doc.add_heading('2. Top 10 Positions', level=1)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
header_cells = table.rows[0].cells
header_cells[0].text = 'Position'
header_cells[1].text = 'Count'

for position, count in top_positions.items():
    row_cells = table.add_row().cells
    row_cells[0].text = str(position)
    row_cells[1].text = str(count)

# 3. Monthly Analysis
df['Month'] = pd.to_datetime(df['Month'])
monthly_counts = df.groupby(df['Month'].dt.strftime('%Y-%m')).size()

# Create a bar plot for monthly counts
plt.figure(figsize=(12, 6))
monthly_counts.plot(kind='bar')
plt.title('Number of Positions by Month')
plt.xlabel('Month')
plt.ylabel('Number of Positions')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('monthly_positions.png')
plt.close()

# Add the plot to the document
doc.add_heading('3. Monthly Analysis', level=1)
doc.add_picture('monthly_positions.png', width=Inches(6))
doc.add_paragraph(f'Peak month: {monthly_counts.idxmax()} with {monthly_counts.max()} positions')

# 4. Location Analysis
location_counts = df['Location'].value_counts()
doc.add_heading('4. Location Analysis', level=1)
for location, count in location_counts.items():
    doc.add_paragraph(f'{location}: {count} positions')

# Save the document
doc.save('Siemens_Staff_Analysis.docx')

print("Analysis complete! Report generated as 'Siemens_Staff_Analysis.docx'")
