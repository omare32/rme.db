import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from datetime import datetime
import numpy as np
import os
import sys

# Enable console output to be captured
sys.stdout.reconfigure(encoding='utf-8')

print("Script started")
print(f"Current working directory: {os.getcwd()}")

try:
    # Read the Excel file
    print("Reading Excel file...")
    excel_path = 'Siemens Wages.xlsx'
    print(f"Excel file path: {os.path.abspath(excel_path)}")
    
    if not os.path.exists(excel_path):
        print(f"ERROR: File {excel_path} does not exist!")
        sys.exit(1)
    
    df = pd.read_excel(excel_path)
    
    # Display basic info about the dataframe
    print("\nDataframe info:")
    print(f"Shape: {df.shape}")
    print("\nColumns:")
    for col in df.columns:
        print(f"- {col}")
    
    print("\nFirst 5 rows of data:")
    print(df.head().to_string())
    
    # Convert date column if it exists
    if 'date' in df.columns:
        print("\nConverting date column...")
        df['date'] = pd.to_datetime(df['date'], errors='coerce')
    
    # Convert numeric columns
    print("\nConverting numeric columns...")
    if 'mandays' in df.columns:
        df['mandays'] = pd.to_numeric(df['mandays'], errors='coerce')
        df['mandays'] = df['mandays'].fillna(0)
    
    if 'amount' in df.columns:
        df['amount'] = pd.to_numeric(df['amount'], errors='coerce')
        df['amount'] = df['amount'].fillna(0)
    
    print("\nColumn types after conversion:")
    print(df.dtypes)
    
    # Create a new Word document
    print("\nCreating Word document...")
    doc = Document()
    doc.add_heading('Siemens Wages Analysis Report', 0)
    doc.add_paragraph(f'Report generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # 1. Monthly Analysis
    print("\nPerforming monthly analysis...")
    if 'date' in df.columns and pd.api.types.is_datetime64_dtype(df['date']):
        monthly_stats = df.groupby(df['date'].dt.strftime('%Y-%m')).agg({
            'mandays': 'sum',
            'amount': 'sum',
        }).reset_index()
        
        print("\nMonthly stats:")
        print(monthly_stats.head().to_string())
        
        # Plot monthly mandays
        print("\nCreating monthly mandays chart...")
        plt.figure(figsize=(12, 6))
        plt.bar(monthly_stats['date'], monthly_stats['mandays'])
        plt.title('Total Mandays per Month')
        plt.xlabel('Month')
        plt.ylabel('Mandays')
        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig('monthly_mandays.png')
        plt.close()
        
        doc.add_heading('1. Monthly Mandays Analysis', level=1)
        doc.add_picture('monthly_mandays.png', width=Inches(6))
        
        # Plot monthly amounts
        print("\nCreating monthly amounts chart...")
        plt.figure(figsize=(12, 6))
        plt.bar(monthly_stats['date'], monthly_stats['amount'])
        plt.title('Total Amount per Month')
        plt.xlabel('Month')
        plt.ylabel('Amount')
        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig('monthly_amounts.png')
        plt.close()
        
        doc.add_heading('2. Monthly Amount Analysis', level=1)
        doc.add_picture('monthly_amounts.png', width=Inches(6))
    else:
        print("WARNING: Date column not found or not properly formatted")
        doc.add_paragraph("Monthly analysis could not be performed due to missing or invalid date data.")
    
    # 2. Occupation Analysis
    print("\nPerforming occupation analysis...")
    if 'occupation' in df.columns:
        top_occupations = df.groupby('occupation').agg({
            'mandays': 'sum',
            'amount': 'sum'
        }).sort_values('amount', ascending=False).head(10)
        
        print("\nTop occupations:")
        print(top_occupations.head().to_string())
        
        # Plot top occupations by amount
        print("\nCreating top occupations by amount chart...")
        plt.figure(figsize=(12, 6))
        sns.barplot(x=top_occupations.index, y=top_occupations['amount'])
        plt.title('Top 10 Occupations by Total Amount')
        plt.xlabel('Occupation')
        plt.ylabel('Total Amount')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        plt.savefig('top_occupations_amount.png')
        plt.close()
        
        doc.add_heading('3. Top Occupations Analysis', level=1)
        doc.add_picture('top_occupations_amount.png', width=Inches(6))
        
        # Plot top occupations by mandays
        print("\nCreating top occupations by mandays chart...")
        plt.figure(figsize=(12, 6))
        sns.barplot(x=top_occupations.index, y=top_occupations['mandays'])
        plt.title('Top 10 Occupations by Total Mandays')
        plt.xlabel('Occupation')
        plt.ylabel('Total Mandays')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        plt.savefig('top_occupations_mandays.png')
        plt.close()
        
        doc.add_picture('top_occupations_mandays.png', width=Inches(6))
    else:
        print("WARNING: Occupation column not found")
        doc.add_paragraph("Occupation analysis could not be performed due to missing occupation data.")
    
    # 3. Statistical Summary
    print("\nCreating statistical summary...")
    doc.add_heading('4. Statistical Summary', level=1)
    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = 'Table Grid'
    stats_table.rows[0].cells[0].text = 'Metric'
    stats_table.rows[0].cells[1].text = 'Value'
    
    stats = [
        ('Total Amount Spent', f"{df['amount'].sum():,.2f}"),
        ('Total Mandays', f"{df['mandays'].sum():,.0f}"),
        ('Average Daily Rate', f"{df['amount'].sum() / max(df['mandays'].sum(), 1):,.2f}"),
    ]
    
    if 'occupation' in df.columns:
        stats.append(('Number of Unique Occupations', f"{df['occupation'].nunique()}"))
    
    if 'date' in df.columns and pd.api.types.is_datetime64_dtype(df['date']) and not monthly_stats.empty:
        stats.append(('Peak Month (by Amount)', f"{monthly_stats.loc[monthly_stats['amount'].idxmax(), 'date']}"))
    
    for stat, value in stats:
        row = stats_table.add_row()
        row.cells[0].text = stat
        row.cells[1].text = str(value)
    
    # 4. Scatter plot of mandays vs amount
    print("\nCreating correlation chart...")
    plt.figure(figsize=(10, 6))
    plt.scatter(df['mandays'], df['amount'])
    plt.title('Correlation between Mandays and Amount')
    plt.xlabel('Mandays')
    plt.ylabel('Amount')
    plt.tight_layout()
    plt.savefig('mandays_amount_correlation.png')
    plt.close()
    
    doc.add_heading('5. Mandays vs Amount Correlation', level=1)
    doc.add_picture('mandays_amount_correlation.png', width=Inches(6))
    
    # 5. Rate Analysis
    print("\nPerforming rate analysis...")
    df['rate'] = df['amount'] / df['mandays'].replace(0, np.nan)
    df['rate'] = df['rate'].fillna(0)
    
    if 'rate' in df.columns and 'occupation' in df.columns:
        # Average rate by occupation
        rate_by_occupation = df.groupby('occupation')['rate'].mean().sort_values(ascending=False).head(10)
        
        plt.figure(figsize=(12, 6))
        sns.barplot(x=rate_by_occupation.index, y=rate_by_occupation.values)
        plt.title('Average Daily Rate by Occupation (Top 10)')
        plt.xlabel('Occupation')
        plt.ylabel('Average Daily Rate')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        plt.savefig('rate_by_occupation.png')
        plt.close()
        
        doc.add_heading('6. Rate Analysis', level=1)
        doc.add_picture('rate_by_occupation.png', width=Inches(6))
    
    # Save the document
    print("\nSaving Word document...")
    doc.save('Siemens_Wages_Analysis.docx')
    
    print("\nAnalysis complete! Report generated as 'Siemens_Wages_Analysis.docx'")

except Exception as e:
    print(f"ERROR: {str(e)}")
    import traceback
    traceback.print_exc()
