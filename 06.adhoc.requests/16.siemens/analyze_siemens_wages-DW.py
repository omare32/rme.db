import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from datetime import datetime
import numpy as np
import os

# Set matplotlib to use a font that supports Arabic
plt.rcParams['font.family'] = 'Arial'
plt.rcParams['font.size'] = 10

# Print current working directory
print(f"Current working directory: {os.getcwd()}")

# Read the Excel file
print("Reading Excel file...")
df = pd.read_excel('Siemens Wages.xlsx')

# Display first few rows to understand the data structure
print("\nFirst 5 rows of data:")
print(df.head())

# Convert numeric columns to appropriate types
print("\nColumn types before conversion:")
print(df.dtypes)

# Convert mandays and amount to numeric, coercing errors to NaN
df['mandays'] = pd.to_numeric(df['mandays'], errors='coerce')
df['amount'] = pd.to_numeric(df['amount'], errors='coerce')

# Fill NaN values with 0
df['mandays'] = df['mandays'].fillna(0)
df['amount'] = df['amount'].fillna(0)

print("\nColumn types after conversion:")
print(df.dtypes)

# Create a new Word document
doc = Document()
doc.add_heading('Siemens Wages Analysis Report', 0)
doc.add_paragraph(f'Report generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

# 1. Monthly Analysis
df['date'] = pd.to_datetime(df['date'])
monthly_stats = df.groupby(df['date'].dt.strftime('%Y-%m')).agg({
    'mandays': 'sum',
    'amount': 'sum',
    'occupation': 'count'
}).reset_index()

# Plot monthly mandays
plt.figure(figsize=(12, 6))
plt.bar(monthly_stats['date'], monthly_stats['mandays'])
plt.title('Total Mandays per Month')
plt.xlabel('Month')
plt.ylabel('Mandays')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('monthly_mandays.png')
plt.close()

doc.add_heading('1. Monthly Mandays Analysis', level=1)
doc.add_picture('monthly_mandays.png', width=Inches(6))

# Plot monthly amounts
plt.figure(figsize=(12, 6))
plt.bar(monthly_stats['date'], monthly_stats['amount'])
plt.title('Total Amount per Month')
plt.xlabel('Month')
plt.ylabel('Amount')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('monthly_amounts.png')
plt.close()

doc.add_heading('2. Monthly Amount Analysis', level=1)
doc.add_picture('monthly_amounts.png', width=Inches(6))

# 2. Occupation Analysis
top_occupations = df.groupby('occupation').agg({
    'mandays': 'sum',
    'amount': 'sum'
}).sort_values('amount', ascending=False).head(10)

# Plot top occupations by amount
plt.figure(figsize=(12, 6))
sns.barplot(data=top_occupations.reset_index(), x='occupation', y='amount')
plt.title('Top 10 Occupations by Total Amount')
plt.xlabel('Occupation')
plt.ylabel('Total Amount')
plt.xticks(rotation=45, ha='right')
plt.tight_layout()
plt.savefig('top_occupations_amount.png')
plt.close()

doc.add_heading('3. Top Occupations Analysis', level=1)
doc.add_picture('top_occupations_amount.png', width=Inches(6))

# Plot top occupations by mandays
plt.figure(figsize=(12, 6))
sns.barplot(data=top_occupations.reset_index(), x='occupation', y='mandays')
plt.title('Top 10 Occupations by Total Mandays')
plt.xlabel('Occupation')
plt.ylabel('Total Mandays')
plt.xticks(rotation=45, ha='right')
plt.tight_layout()
plt.savefig('top_occupations_mandays.png')
plt.close()

doc.add_picture('top_occupations_mandays.png', width=Inches(6))

# 3. Statistical Summary
doc.add_heading('4. Statistical Summary', level=1)
stats_table = doc.add_table(rows=1, cols=2)
stats_table.style = 'Table Grid'
stats_table.rows[0].cells[0].text = 'Metric'
stats_table.rows[0].cells[1].text = 'Value'

stats = [
    ('Total Amount Spent', f"{df['amount'].sum():,.2f}"),
    ('Total Mandays', f"{df['mandays'].sum():,.0f}"),
    ('Average Daily Rate', f"{df['amount'].sum() / df['mandays'].sum():,.2f}"),
    ('Number of Unique Occupations', f"{df['occupation'].nunique()}"),
    ('Peak Month (by Amount)', f"{monthly_stats.loc[monthly_stats['amount'].idxmax(), 'date']}")
]

for stat, value in stats:
    row = stats_table.add_row()
    row.cells[0].text = stat
    row.cells[1].text = str(value)

# 4. Scatter plot of mandays vs amount
plt.figure(figsize=(10, 6))
plt.scatter(df['mandays'], df['amount'])
plt.title('Correlation between Mandays and Amount')
plt.xlabel('Mandays')
plt.ylabel('Amount')
plt.tight_layout()
plt.savefig('mandays_amount_correlation.png')
plt.close()

doc.add_heading('5. Mandays vs Amount Correlation', level=1)
doc.add_picture('mandays_amount_correlation.png', width=Inches(6))

# Save the document
doc.save('Siemens_Wages_Analysis.docx')

print("Analysis complete! Report generated as 'Siemens_Wages_Analysis.docx'")
