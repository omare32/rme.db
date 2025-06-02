from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from airflow.models import DagBag, DagRun
from airflow.utils.state import State
from sqlalchemy import desc
from airflow.utils.session import create_session
import os

def get_dag_status():
    """Get status of all DAGs"""
    dag_bag = DagBag()
    dag_status = []
    
    with create_session() as session:
        for dag_id, dag in dag_bag.dags.items():
            # Get the latest run
            latest_run = (
                session.query(DagRun)
                .filter(DagRun.dag_id == dag_id)
                .order_by(desc(DagRun.start_date))
                .first()
            )
            
            status = {
                'dag_id': dag_id,
                'is_active': dag.get_is_active(),
                'schedule_interval': str(dag.schedule_interval),
                'last_run': latest_run.start_date if latest_run else None,
                'last_run_state': latest_run.state if latest_run else None,
                'last_run_duration': (latest_run.end_date - latest_run.start_date).total_seconds() if latest_run and latest_run.end_date else None,
                'description': dag.description or 'No description available'
            }
            dag_status.append(status)
    
    return dag_status

def create_report():
    """Create a Word document with DAG status report"""
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Airflow DAG Status Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add summary section
    doc.add_heading('Summary', level=1)
    dag_status = get_dag_status()
    active_dags = sum(1 for dag in dag_status if dag['is_active'])
    inactive_dags = sum(1 for dag in dag_status if not dag['is_active'])
    
    summary = doc.add_paragraph()
    summary.add_run(f'Total DAGs: {len(dag_status)}\n')
    summary.add_run(f'Active DAGs: {active_dags}\n')
    summary.add_run(f'Inactive DAGs: {inactive_dags}\n')
    
    # Add detailed DAG status
    doc.add_heading('Detailed DAG Status', level=1)
    
    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    headers = ['DAG ID', 'Status', 'Schedule', 'Last Run', 'Last Run State', 'Duration (s)']
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    # Add DAG information
    for dag in dag_status:
        row_cells = table.add_row().cells
        row_cells[0].text = dag['dag_id']
        row_cells[1].text = 'Active' if dag['is_active'] else 'Inactive'
        row_cells[2].text = dag['schedule_interval']
        row_cells[3].text = dag['last_run'].strftime('%Y-%m-%d %H:%M:%S') if dag['last_run'] else 'Never'
        row_cells[4].text = dag['last_run_state'] if dag['last_run_state'] else 'N/A'
        row_cells[5].text = f"{dag['last_run_duration']:.2f}" if dag['last_run_duration'] else 'N/A'
        
        # Color coding for status
        if dag['last_run_state'] == 'success':
            row_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif dag['last_run_state'] == 'failed':
            row_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    # Add description section
    doc.add_heading('DAG Descriptions', level=1)
    for dag in dag_status:
        doc.add_heading(dag['dag_id'], level=2)
        doc.add_paragraph(dag['description'])
    
    # Save the document
    report_path = os.path.join(os.path.dirname(__file__), 'dag_status_report.docx')
    doc.save(report_path)
    return report_path

if __name__ == "__main__":
    report_path = create_report()
    print(f"Report generated successfully at: {report_path}") 