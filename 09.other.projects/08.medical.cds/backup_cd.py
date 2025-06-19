import os
import shutil
import datetime
import json
import hashlib
import win32com.client # For .doc files via Word
import docx          # For .docx files
import PyPDF2        # For .pdf files
# DICOM libraries will be imported lazily in the extraction function
import subprocess

CD_DRIVE = "E:\\"
BACKUP_BASE_DIR = "F:\\My Drive\\Backups\\medical-cds"
# Directory for all text summaries and future AI-generated summaries
SUMMARY_BASE_DIR = os.path.join(BACKUP_BASE_DIR, "summaries")
# Directory to store extracted DICOM images
IMAGES_BASE_DIR = os.path.join(BACKUP_BASE_DIR, "images")

def get_drive_label(drive_path):
    """Gets the volume label of a drive."""
    try:
        # For Windows, using vol command
        import subprocess
        result = subprocess.check_output(['cmd', '/c', f'vol {drive_path}'], universal_newlines=True)
        label = result.split('\n')[0].split(' is ')[-1].strip()
        if not label or "No Volume Label" in label or "Incorrect Volume Label" in label:
            return None
        return label
    except Exception as e:
        print(f"Error getting drive label for {drive_path}: {e}")
        return None

# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def sanitize_label(label: str) -> str:
    """Return a filesystem-safe version of the CD label used in folder names."""
    return "".join(c for c in label if c.isalnum() or c in (' ', '_')).rstrip()


def find_existing_backup_path(cd_label: str):
    """If a backup folder for this CD label already exists under BACKUP_BASE_DIR,
    return its absolute path, else None. We match by the trailing portion
    "_<sanitized_label>" of the folder name.
    """
    if not cd_label:
        return None
    sanitized_label = sanitize_label(cd_label)
    try:
        matches = [d for d in os.listdir(BACKUP_BASE_DIR)
                   if d.endswith(f"_{sanitized_label}") and os.path.isdir(os.path.join(BACKUP_BASE_DIR, d))]
        if matches:
            # If multiple, choose the latest by timestamp (folder name prefix)
            matches.sort(reverse=True)
            return os.path.join(BACKUP_BASE_DIR, matches[0])
    except FileNotFoundError:
        pass
    return None


def create_backup_folder_name(cd_label):
    """Creates a unique folder name for the backup."""
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    if cd_label:
        # Sanitize label to be a valid folder name
        sanitized_label = "".join(c for c in cd_label if c.isalnum() or c in (' ', '_')).rstrip()
        return f"{timestamp}_{sanitized_label}"
    return timestamp

def calculate_sha256(file_path):
    """Calculates the SHA256 checksum of a file."""
    sha256_hash = hashlib.sha256()
    try:
        with open(file_path, "rb") as f:
            # Read and update hash string value in blocks of 4K
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    except FileNotFoundError:
        print(f"Error: File not found at {file_path} during checksum calculation.")
        return None
    except Exception as e:
        print(f"Error calculating checksum for {file_path}: {e}")
        return None

def generate_report(backup_path, source_path, cd_label, copied_items_details, copy_errors, overall_checksum_match, text_summary_file_path=None):
    """Generates a JSON and Markdown report of the backup, including checksums."""
    report_data = {
        "backup_timestamp": datetime.datetime.now().isoformat(),
        "source_drive": source_path,
        "cd_label": cd_label,
        "backup_location": backup_path,
        "overall_status": "",
        "files_and_folders": copied_items_details, 
        "total_items_processed": len(copied_items_details),
        "total_size_bytes": sum(item.get("size_bytes", 0) for item in copied_items_details if item["type"] == "file"),
        "errors_during_copy": copy_errors,
        "all_checksums_match": overall_checksum_match,
        "text_summary_file": text_summary_file_path
    }

    if not copy_errors and overall_checksum_match:
        report_data["overall_status"] = "Success: All files copied and checksums match."
    elif not copy_errors and not overall_checksum_match:
        report_data["overall_status"] = "Warning: Files copied, but one or more checksums mismatched."
    else:
        report_data["overall_status"] = "Error: Backup completed with errors or checksum mismatches."


    markdown_report_lines = [
        f"# Medical CD Backup Report",
        f"- **Backup Date:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"- **Source CD Drive:** {source_path}",
        f"- **CD Label:** {cd_label if cd_label else 'N/A'}",
        f"- **Backup Location:** `{backup_path}`",
        f"- **Overall Status:** {report_data['overall_status']}",
        f"## Backup Details:"
    ]
    
    num_files = 0
    num_dirs = 0

    for item in copied_items_details:
        if item["type"] == "directory":
            num_dirs += 1
            markdown_report_lines.append(f"  - [D] {item['path']}")
        elif item["type"] == "file":
            num_files +=1
            size_mb = item.get('size_bytes', 0) / (1024*1024)
            checksum_status = "Match" if item.get("checksum_match") else "MISMATCH"
            if item.get("source_checksum_sha256") is None or item.get("destination_checksum_sha256") is None:
                 checksum_status = "Error/Unavailable"
            
            line = f"  - [F] {item['path']} ({size_mb:.2f} MB)"
            # Only show checksum info if source_checksum_sha256 is not None (i.e., checksum was attempted)
            if item.get("source_checksum_sha256") is not None:
                 line += f" - SHA256: {checksum_status}"
                 # Show truncated hashes only on mismatch and if both hashes are available
                 if not item.get("checksum_match") and item.get("source_checksum_sha256") and item.get("destination_checksum_sha256"):
                     line += f" (Source: {item['source_checksum_sha256'][:8]}..., Dest: {item['destination_checksum_sha256'][:8]}...)"
            markdown_report_lines.append(line)

    markdown_report_lines.append(f"\n### Summary:")
    markdown_report_lines.append(f"- **Total Directories Processed:** {num_dirs}")
    markdown_report_lines.append(f"- **Total Files Processed:** {num_files}")
    markdown_report_lines.append(f"- **Total Backup Size:** {report_data['total_size_bytes'] / (1024*1024):.2f} MB")

    if text_summary_file_path:
        relative_summary_path = os.path.basename(text_summary_file_path)
        markdown_report_lines.append(f"- **Extracted Text Summary:** [{relative_summary_path}](./{relative_summary_path})")
    
    if copy_errors:
        markdown_report_lines.append(f"\n### Errors Encountered During Copy/Verification:")
        for err in copy_errors:
            markdown_report_lines.append(f"  - {err}")
    elif not overall_checksum_match and any(item['type'] == 'file' and not item.get('checksum_match', True) for item in copied_items_details):
        # This condition ensures we only print mismatch details if there are actual file mismatches, not just overall_checksum_match being false due to other errors
        markdown_report_lines.append(f"\n### Checksum Mismatches Detected:")
        markdown_report_lines.append("  At least one file had a checksum mismatch. Please review file details above.")

    # Save JSON report
    json_report_path = os.path.join(backup_path, "backup_report.json")
    with open(json_report_path, 'w') as f:
        json.dump(report_data, f, indent=4)

    # Save Markdown report
    md_report_path = os.path.join(backup_path, "backup_report.md")
    with open(md_report_path, 'w') as f:
        f.write("\n".join(markdown_report_lines))
    
    print(f"Reports generated: {json_report_path} and {md_report_path}")

def backup_cd():
    """Main function to backup the CD, including checksum verification."""
    if not os.path.exists(CD_DRIVE):
        print(f"Error: CD Drive {CD_DRIVE} not found. Please check the drive letter.")
        return

    try:
        if not os.listdir(CD_DRIVE):
            print(f"Warning: CD Drive {CD_DRIVE} appears to be empty. Consider if this is expected.")
            # Proceeding to allow backup of an 'empty' CD if that's intended.
    except OSError as e:
        print(f"Error accessing CD Drive {CD_DRIVE}: {e}. Please ensure a disc is inserted and ready.")
        return

    print(f"Scanning CD in drive {CD_DRIVE}...")
    cd_label = get_drive_label(CD_DRIVE.rstrip('\\'))
    # Determine scan date from original CD contents before any copy
    pre_copy_scan_date = determine_scan_date(CD_DRIVE)
    print(f"CD Label: {cd_label if cd_label else 'No label found'}")

    # ------------------------------------------------------------------
    # Check for existing backup of this CD to avoid duplicates
    # ------------------------------------------------------------------
    existing_backup_path = find_existing_backup_path(cd_label) if cd_label else None
    if existing_backup_path:
        print(f"Detected an existing backup for label '{cd_label}'. Skipping copy and using folder:\n  {existing_backup_path}")
        # We still attempt text extraction (in case it was not done yet) and then exit.
        text_summary_path = extract_text_from_documents(existing_backup_path)
        if text_summary_path:
            print("Text extraction finished. Summary doc created.")
        else:
            print("No documents found or summary already existed.")

        # ---------------- DICOM image extraction -----------------
        images_folder = extract_images_from_dicom(existing_backup_path)

        # ---------------- Scan date metadata file -----------------
        scan_date = pre_copy_scan_date or determine_scan_date(existing_backup_path, cd_label)
        if scan_date:
            create_scan_date_file(scan_date, backup_folder_name)
        if images_folder:
            print(f"Images extracted to {images_folder}")
        else:
            print("No DICOM images found to extract.")
        return

    # ------------------------------------------------------------------
    # No existing backup; proceed with normal copy procedure
    # ------------------------------------------------------------------

    backup_folder_name = create_backup_folder_name(cd_label)
    destination_path_base = os.path.join(BACKUP_BASE_DIR, backup_folder_name)

    copied_items_details = [] 
    errors_during_copy = []
    overall_checksum_match = True # Assume true until a mismatch or error occurs

    try:
        if not os.path.exists(BACKUP_BASE_DIR):
            os.makedirs(BACKUP_BASE_DIR)
            print(f"Created base backup directory: {BACKUP_BASE_DIR}")

        if os.path.exists(destination_path_base):
            print(f"Error: Destination folder {destination_path_base} already exists. Aborting to prevent overwrite.")
            return
        
        os.makedirs(destination_path_base)
        print(f"Created backup directory: {destination_path_base}")
        print(f"Starting backup of {CD_DRIVE} to {destination_path_base} with checksum verification...")

        for root, dirs, files in os.walk(CD_DRIVE):
            # Make sure relative_path handles the root of the CD drive correctly ('.')
            relative_path_from_cd_root = os.path.relpath(root, CD_DRIVE)
            if relative_path_from_cd_root == '.':
                destination_root = destination_path_base
                current_relative_path_for_report = '' # For root items
            else:
                destination_root = os.path.join(destination_path_base, relative_path_from_cd_root)
                current_relative_path_for_report = relative_path_from_cd_root

            for dir_name in dirs:
                # source_dir_path = os.path.join(root, dir_name) # Not needed for makedirs
                dest_dir_path = os.path.join(destination_root, dir_name)
                try:
                    os.makedirs(dest_dir_path, exist_ok=True)
                    copied_items_details.append({
                        "type": "directory",
                        "path": os.path.join(current_relative_path_for_report, dir_name)
                    })
                except Exception as e:
                    err_msg = f"Error creating directory {dest_dir_path}: {e}"
                    print(err_msg)
                    errors_during_copy.append(err_msg)

            for file_name in files:
                source_file_path = os.path.join(root, file_name)
                dest_file_path = os.path.join(destination_root, file_name)
                
                # Ensure parent directory for the file exists in destination
                os.makedirs(os.path.dirname(dest_file_path), exist_ok=True)

                file_detail = {
                    "type": "file",
                    "path": os.path.join(current_relative_path_for_report, file_name),
                    "size_bytes": 0,
                    "source_checksum_sha256": None,
                    "destination_checksum_sha256": None,
                    "checksum_match": False # Default to false, prove true
                }

                try:
                    file_detail["size_bytes"] = os.path.getsize(source_file_path)
                    
                    print(f"Copying {file_name}...") # Simplified print
                    shutil.copy2(source_file_path, dest_file_path)

                    print(f"Verifying {file_name}...")
                    file_detail["source_checksum_sha256"] = calculate_sha256(source_file_path)
                    file_detail["destination_checksum_sha256"] = calculate_sha256(dest_file_path)

                    if file_detail["source_checksum_sha256"] and \
                       file_detail["source_checksum_sha256"] == file_detail["destination_checksum_sha256"]:
                        file_detail["checksum_match"] = True
                    else:
                        file_detail["checksum_match"] = False
                        overall_checksum_match = False # A single mismatch makes the overall false
                        # Log specific mismatch if source or dest checksum is None (error during calculation)
                        if file_detail["source_checksum_sha256"] is None or file_detail["destination_checksum_sha256"] is None:
                            err_msg = f"Checksum verification failed for {file_name} due to error in hash calculation."
                        else:
                            err_msg = f"Checksum MISMATCH for {file_name}!"
                        print(err_msg)
                        errors_during_copy.append(err_msg) # Add to general errors or a specific list
                
                except FileNotFoundError as fnf_err:
                    err_msg = f"Error: File not found during copy/verification of {source_file_path}: {fnf_err}"
                    print(err_msg)
                    errors_during_copy.append(err_msg)
                    overall_checksum_match = False 
                except Exception as e:
                    err_msg = f"Error copying/verifying file {source_file_path}: {e}"
                    print(err_msg)
                    errors_during_copy.append(err_msg)
                    overall_checksum_match = False 
                
                copied_items_details.append(file_detail)

        if not errors_during_copy and overall_checksum_match:
            print("Backup completed successfully. All checksums match.")
        elif not errors_during_copy and not overall_checksum_match:
            print("Backup completed. IMPORTANT: One or more file checksums MISMATCHED. Check report for details.")
        else:
            print("Backup completed with errors or checksum mismatches. Check report for details.")

        # Attempt to extract text from documents
        text_summary_path = None
        try:
            text_summary_path = extract_text_from_documents(destination_path_base)
            images_folder = extract_images_from_dicom(destination_path_base)
            scan_date = pre_copy_scan_date or determine_scan_date(destination_path_base, cd_label)
            if scan_date:
                create_scan_date_file(scan_date, backup_folder_name)

        except ImportError:
            print("\nSkipping text extraction: required libraries not available. Ensure MS Word, python-docx, and PyPDF2 are installed.")
            errors_during_copy.append("Text extraction skipped: missing dependencies.")
        except Exception as e_textract:
            print(f"\nError during text extraction process: {e_textract}")
            errors_during_copy.append(f"Text extraction failed: {e_textract}")

        print("\nGenerating final backup report...")
        generate_report(destination_path_base, CD_DRIVE, cd_label, copied_items_details, errors_during_copy, overall_checksum_match, text_summary_path)
        print(f"Backup, text extraction (if any), and report generation finished for CD '{cd_label if cd_label else 'Unnamed CD'}' dated {backup_folder_name.split('_')[0]}. ")
        print(f"You can find the backup in: {destination_path_base}")

    except FileNotFoundError as e:
        print(f"Critical Error: {e}. This might indicate the CD drive {CD_DRIVE} became unavailable or a path issue.")
    except Exception as e:
        print(f"An unexpected critical error occurred during the backup process: {e}")
        # Consider if destination_path_base should be cleaned up or left for inspection
        # if os.path.exists(destination_path_base):
        #     print(f"Cleaning up partially created folder: {destination_path_base}")
        #     shutil.rmtree(destination_path_base)

def determine_scan_date(backup_folder_path, cd_label=None, backup_folder_name=None):
    """Determine scan date from all possible sources and return the earliest as YYYYMMDD."""
    import datetime
    possible_dates = set()
    try:
        import pydicom
    except ImportError:
        pydicom = None

    # 1. DICOM metadata
    if pydicom:
        for root, _, files in os.walk(backup_folder_path):
            for fname in files:
                if fname.lower().endswith('.dcm'):
                    fpath = os.path.join(root, fname)
                    try:
                        ds = pydicom.dcmread(fpath, stop_before_pixels=True)
                        for tag in ("StudyDate", "SeriesDate", "AcquisitionDate", "ContentDate"):
                            if tag in ds and ds.get(tag):
                                val = str(ds.get(tag).value).strip()
                                if val and len(val) == 8 and val.isdigit():
                                    possible_dates.add(val)
                    except Exception:
                        continue
    # 2. All file mtimes
    for root, _, files in os.walk(backup_folder_path):
        for fname in files:
            fpath = os.path.join(root, fname)
            try:
                ts = os.path.getmtime(fpath)
                dt = datetime.datetime.fromtimestamp(ts).strftime("%Y%m%d")
                possible_dates.add(dt)
            except Exception:
                continue
    # 3. Parse date from CD label
    if cd_label:
        import re
        m = re.search(r'(\d{2})(\d{2})(\d{4})', cd_label)
        if m:
            possible_dates.add(f"{m.group(3)}{m.group(2)}{m.group(1)}")
    # 4. Parse date from backup folder name (if matches pattern)
    if backup_folder_name:
        import re
        m = re.search(r'(\d{8})', backup_folder_name)
        if m:
            possible_dates.add(m.group(1))
    # Return the earliest valid date
    if possible_dates:
        return min(possible_dates)
    return None


def create_scan_date_file(scan_date, backup_folder_name):
    """Create a single text file named 'Data YYYY.MM.DD.txt' in the summaries folder."""
    import os
    # Format YYYYMMDD -> YYYY.MM.DD
    if len(scan_date) == 8 and scan_date.isdigit():
        dotted = f"{scan_date[0:4]}.{scan_date[4:6]}.{scan_date[6:8]}"
    else:
        dotted = scan_date
    fname = f"Data {dotted}.txt"
    summaries_dir = os.path.join(BACKUP_BASE_DIR, 'summaries')
    if not os.path.exists(summaries_dir):
        try:
            os.makedirs(summaries_dir)
            print(f"[DEBUG] Created summaries directory: {summaries_dir}")
        except Exception as e:
            print(f"Could not create summaries directory: {e}")
            return
    date_file_path = os.path.join(summaries_dir, fname)
    print(f"[DEBUG] Attempting to create scan date file at: {date_file_path}")
    if not os.path.exists(date_file_path):
        try:
            with open(date_file_path, 'w') as f:
                f.write(f"Scan date (from metadata, file timestamps, or CD label): {dotted}\n")
            print(f"[DEBUG] Successfully wrote scan date file: {date_file_path}")
        except Exception as e:
            print(f"Could not write scan date file: {e}")


def extract_images_from_dicom(backup_folder_path):
    """Extract images from DICOM files under a backup folder.
    Saves PNG files into IMAGES_BASE_DIR/<backup_folder_name>/
    Returns the images folder path or None if none extracted.
    """
    try:
        import pydicom
        import numpy as np
        from PIL import Image
    except ImportError:
        print("pydicom / numpy / pillow not available, skipping DICOM image extraction.")
        return None

    os.makedirs(IMAGES_BASE_DIR, exist_ok=True)
    backup_folder_name = os.path.basename(backup_folder_path.rstrip(os.sep))
    output_dir = os.path.join(IMAGES_BASE_DIR, backup_folder_name)
    os.makedirs(output_dir, exist_ok=True)

    count = 0
    for root, _, files in os.walk(backup_folder_path):
        for fname in files:
            # We attempt to read any file; pydicom will skip non-DICOM gracefully.
            if not fname.lower().endswith('.dcm'):
                continue
                # attempt to open any file; pydicom will raise if not dicom
                pass  # we'll still try
            file_path = os.path.join(root, fname)
            try:
                ds = pydicom.dcmread(file_path, force=True, stop_before_pixels=False)
                if hasattr(ds, 'pixel_array'):
                    arr = ds.pixel_array
                    # Rescale to 8-bit if necessary
                    if arr.dtype != np.uint8:
                        arr = arr.astype(np.float32)
                        arr = 255 * (arr - arr.min()) / (arr.max() - arr.min() + 1e-5)
                        arr = arr.astype(np.uint8)
                    img = Image.fromarray(arr)
                    out_name = f"{count:04d}.png"
                    img.save(os.path.join(output_dir, out_name))
                    count += 1
            except Exception:
                continue

    return output_dir if count else None


def extract_text_from_documents(backup_folder_path):
    """Extracts text from DOC, DOCX, and PDF files found in a CD backup folder and
    saves the combined content into a Word (.docx) file located under
    `SUMMARY_BASE_DIR`. The summary filename is based on the backup folder name.

    Returns the full path to the created summary document, or None if no
    documents were found / processed.
    """
    supported_extensions = ('.doc', '.docx', '.pdf')
    extracted_entries = []  # list of tuples (relative_path, text)
    processed_files_count = 0

    # Ensure summary directory exists
    os.makedirs(SUMMARY_BASE_DIR, exist_ok=True)

    backup_folder_name = os.path.basename(backup_folder_path.rstrip(os.sep))
    summary_file_name = f"{backup_folder_name}_text_summary.docx"
    summary_file_path = os.path.join(SUMMARY_BASE_DIR, summary_file_name)

    print(f"\nScanning for documents to summarize in: {backup_folder_path}")

    # Initialize Word COM object if .doc files might be present
    word_app = None
    try:
        has_doc_files = any(f.lower().endswith('.doc') for r, d, fs in os.walk(backup_folder_path) for f in fs)
        if has_doc_files:
            print("  Initializing MS Word for .doc file processing (this may take a moment)...")
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
    except Exception as e:
        print(f"    Warning: Could not initialize MS Word COM object. .doc file extraction will be skipped. Error: {e}")
        word_app = None

    # Walk the backup folder and extract text file-by-file
    for root, _, files in os.walk(backup_folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            relative_file_path = os.path.relpath(file_path, backup_folder_path)
            lower_name = file.lower()
            text = ""

            if lower_name.endswith('.doc'):
                if word_app:
                    print(f"  Processing .doc: {relative_file_path}")
                    try:
                        doc = word_app.Documents.Open(file_path, ReadOnly=True)
                        text = doc.Content.Text
                        doc.Close(False)
                        processed_files_count += 1
                    except Exception as e:
                        print(f"    Error extracting text from .doc {relative_file_path}: {e}")
                        text = f"[Error extracting text via Word: {e}]"
                else:
                    print(f"    Attempting antiword fallback for .doc: {relative_file_path}")
                    try:
                        result = subprocess.run(["antiword", file_path], capture_output=True, text=True, timeout=30)
                        if result.returncode == 0:
                            text = result.stdout
                            processed_files_count += 1
                        else:
                            text = f"[antiword failed: {result.stderr.strip() or 'unknown error'}]"
                    except FileNotFoundError:
                        text = "[antiword executable not found in PATH]"
                    except Exception as e2:
                        text = f"[antiword error: {e2}]"

            elif lower_name.endswith('.docx'):
                print(f"  Processing .docx: {relative_file_path}")
                try:
                    d = docx.Document(file_path)
                    text = '\n'.join(p.text for p in d.paragraphs)
                    processed_files_count += 1
                except Exception as e:
                    print(f"    Error extracting text from .docx {relative_file_path}: {e}")
                    text = f"[Error extracting text: {e}]"

            elif lower_name.endswith('.pdf'):
                print(f"  Processing .pdf: {relative_file_path}")
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        page_text = []
                        for page in reader.pages:
                            page_text.append(page.extract_text())
                    text = '\n'.join(filter(None, page_text))
                    processed_files_count += 1
                except Exception as e:
                    print(f"    Error extracting text from .pdf {relative_file_path}: {e}")
                    text = f"[Error extracting text: {e}]"

            if text:
                extracted_entries.append((relative_file_path, text))

    # Close Word COM if opened
    if word_app:
        try:
            word_app.Quit(False)
        except Exception:
            pass

    if not extracted_entries:
        print("No documents found or no text extracted for summary.")
        return None

    # Write to Word document using python-docx
    try:
        from docx import Document  # local import to avoid confusion with var name
        doc_summary = Document()
        doc_summary.add_heading("Extracted Text from Documents", level=1)

        for rel_path, entry_text in extracted_entries:
            doc_summary.add_heading(f"Content from: {rel_path}", level=2)
            for line in entry_text.split('\n'):
                doc_summary.add_paragraph(line)
            doc_summary.add_page_break()

        doc_summary.save(summary_file_path)
        print(f"\nWord summary created: {summary_file_path} (from {processed_files_count} document(s))")
        return summary_file_path
    except Exception as e:
        print(f"Error writing Word summary file: {e}")
        return None

    summary_file_name = "text_summary.md"
    summary_file_path = os.path.join(backup_folder_path, summary_file_name)
    processed_files_count = 0

    print(f"\nScanning for documents to summarize in: {backup_folder_path}")

    # Initialize Word COM object if .doc files might be present
    word_app = None
    try:
        # Check if any .doc files exist to avoid starting Word unnecessarily
        has_doc_files = any(f.lower().endswith('.doc') for r, d, fs in os.walk(backup_folder_path) for f in fs)
        if has_doc_files:
            print("  Initializing MS Word for .doc file processing (this may take a moment)...")
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False # Run Word in the background
    except Exception as e:
        print(f"    Warning: Could not initialize MS Word COM object. .doc file extraction will be skipped. Error: {e}")
        word_app = None # Ensure it's None if initialization failed

    for root, _, files in os.walk(backup_folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            relative_file_path = os.path.relpath(file_path, backup_folder_path)
            text = ""
            extracted = False

            if file.lower().endswith('.doc'):
                if word_app:
                    print(f"  Processing .doc: {relative_file_path}")
                    try:
                        doc = word_app.Documents.Open(file_path, ReadOnly=True)
                        text = doc.Content.Text
                        doc.Close(False) # Close without saving changes
                        extracted = True
                    except Exception as e:
                        print(f"    Error extracting text from .doc file {relative_file_path} using Word: {e}")
                        text = f"[Error extracting text via Word: {e}]"
                else:
                    print(f"    Skipping .doc file {relative_file_path} as MS Word COM object is not available.")
                    text = "[Skipped .doc file - MS Word COM not available]"
            
            elif file.lower().endswith('.docx'):
                print(f"  Processing .docx: {relative_file_path}")
                try:
                    doc = docx.Document(file_path)
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    text = '\n'.join(full_text)
                    extracted = True
                except Exception as e:
                    print(f"    Error extracting text from .docx file {relative_file_path}: {e}")
                    text = f"[Error extracting text: {e}]"

            elif file.lower().endswith('.pdf'):
                print(f"  Processing .pdf: {relative_file_path}")
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        full_text = []
                        for page_num in range(len(reader.pages)):
                            page = reader.pages[page_num]
                            full_text.append(page.extract_text())
                        text = '\n'.join(filter(None, full_text)) # Filter out None results from extract_text
                    extracted = True
                except Exception as e:
                    print(f"    Error extracting text from .pdf file {relative_file_path}: {e}")
                    text = f"[Error extracting text: {e}]"
            
            if extracted or text: # If extracted is true or text has some error message
                extracted_texts.append(f"## Content from: {relative_file_path}\n\n{text}\n\n---\n")
                if extracted: processed_files_count += 1

    if word_app:
        try:
            word_app.Quit(False) # Quit Word without saving changes
            print("  MS Word application closed.")
        except Exception as e:
            print(f"    Warning: Error trying to quit MS Word: {e}")
    
    if extracted_texts:
        try:
            with open(summary_file_path, 'w', encoding='utf-8') as f:
                f.write("# Extracted Text from Documents\n\n")
                f.write("\n".join(extracted_texts))
            print(f"\nText summary from {processed_files_count} document(s) saved to: {summary_file_path}")
            return summary_file_path
        except Exception as e:
            print(f"Error writing text summary file: {e}")
            return None
    else:
        print("No documents found or no text successfully extracted for summary.")
        return None


if __name__ == "__main__":
    backup_cd()
