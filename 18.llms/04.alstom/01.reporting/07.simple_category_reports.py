import os
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_summaries():
    """Load existing summaries if available"""
    data_dir = os.path.join(os.path.dirname(__file__), "data")
    summaries_path = os.path.join(data_dir, "ollama_summaries.json")
    
    if not os.path.exists(summaries_path):
        print(f"No summaries found at {summaries_path}")
        return {}
    
    try:
        with open(summaries_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Error loading summaries: {str(e)}")
        return {}

def load_pdf_files():
    """Get all PDF files from the Alstom directory"""
    base_paths = ["C:\\alstom\\folder1", "C:\\alstom\\folder2"]
    pdf_files = []
    
    for base_path in base_paths:
        for root, _, files in os.walk(base_path):
            for file in files:
                if file.lower().endswith('.pdf'):
                    full_path = os.path.join(root, file)
                    simplified_path = os.path.relpath(os.path.dirname(full_path), base_path)
                    pdf_files.append({
                        'name': file,
                        'full_path': full_path,
                        'simplified_path': simplified_path
                    })
    
    print(f"Found {len(pdf_files)} PDF files in total")
    return pdf_files

def categorize_documents(pdf_files, summaries):
    """Categorize documents based on keywords in their names and paths"""
    categories = {
        "Technical Specifications": [],
        "Design Drawings": [],
        "Project Management": [],
        "Contracts and Agreements": [],
        "Reports and Assessments": [],
        "Safety and Compliance": [],
        "Miscellaneous": []
    }
    
    # Define keywords for each category
    category_keywords = {
        "Technical Specifications": ["specification", "spec", "technical", "requirement", "standard"],
        "Design Drawings": ["drawing", "design", "plan", "layout", "sketch", "diagram", "ifc"],
        "Project Management": ["schedule", "timeline", "milestone", "progress", "management", "plan"],
        "Contracts and Agreements": ["contract", "agreement", "legal", "terms", "condition", "scope"],
        "Reports and Assessments": ["report", "assessment", "analysis", "evaluation", "survey"],
        "Safety and Compliance": ["safety", "compliance", "regulation", "standard", "quality", "inspection"],
        "Miscellaneous": []
    }
    
    # Assign each document to a category
    for doc in pdf_files:
        assigned = False
        doc_text = doc['name'].lower() + ' ' + doc['simplified_path'].lower()
        
        # Add summary to doc_text if available
        if doc['name'] in summaries:
            doc['summary'] = summaries[doc['name']]
            doc_text += ' ' + summaries[doc['name']].lower()
        
        # Check each category
        for category, keywords in category_keywords.items():
            for keyword in keywords:
                if keyword in doc_text:
                    categories[category].append(doc)
                    assigned = True
                    break
            
            if assigned:
                break
        
        # If not assigned to any category, put in Miscellaneous
        if not assigned:
            categories["Miscellaneous"].append(doc)
    
    return categories

def create_category_report(category_name, docs, output_path):
    """Create a Word document with summaries for a specific category"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Alstom Project Documentation: {category_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    intro = doc.add_paragraph()
    intro.add_run('Executive Summary').bold = True
    doc.add_paragraph(f'This report contains summaries of important {category_name.lower()} documents from the Alstom project. These documents have been selected based on their relevance to key project aspects and organized to provide a comprehensive overview.')
    
    doc.add_paragraph()  # Add space
    
    # Add document summaries
    doc.add_heading('Document Summaries', 1)
    
    for i, doc_info in enumerate(docs, 1):
        name = doc_info['name']
        path = doc_info['simplified_path']
        
        # Add document header with formatting
        header = doc.add_paragraph()
        header.add_run(f"{i}. {name}").bold = True
        header.add_run(f"\nLocation: {path}")
        
        # Add summary with formatting
        if 'summary' in doc_info:
            summary = doc_info['summary']
            p = doc.add_paragraph()
            p.add_run('Summary: ').bold = True
            p.add_run(summary)
        else:
            doc.add_paragraph('Summary not available for this document.').italic = True
            
        doc.add_paragraph()  # Add space between documents
    
    # Save the document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\nReport saved to: {output_path}")

def create_master_report(categorized_docs, output_path):
    """Create a master Word document with all documents organized by category"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Alstom Project Documentation Master Index', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    intro = doc.add_paragraph()
    intro.add_run('Executive Summary').bold = True
    doc.add_paragraph('This report provides a comprehensive index of all important documents from the Alstom project, organized by category. Each category has its own detailed report with full summaries.')
    
    doc.add_paragraph()  # Add space
    
    # Add table of contents header
    doc.add_heading('Document Categories', 1)
    
    # Add each category
    for category, docs in categorized_docs.items():
        if docs:  # Only add category if it has documents
            category_heading = doc.add_heading(category, 2)
            
            # Add document list
            for i, doc_info in enumerate(docs, 1):
                name = doc_info['name']
                path = doc_info['simplified_path']
                
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{name}").bold = True
                p.add_run(f" (in {path})")
            
            doc.add_paragraph()  # Add space between categories
    
    # Save the document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\nMaster index saved to: {output_path}")

def main():
    # Load summaries
    print("Loading summaries...")
    summaries = load_summaries()
    print(f"Loaded {len(summaries)} summaries.")
    
    # Load PDF files
    print("\nLoading PDF files...")
    pdf_files = load_pdf_files()
    
    # Categorize documents
    print("\nCategorizing documents...")
    categorized_docs = categorize_documents(pdf_files, summaries)
    
    # Create reports directory
    reports_dir = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(reports_dir, exist_ok=True)
    
    # Create category reports
    print("\nGenerating category reports...")
    for category, docs in categorized_docs.items():
        if docs:  # Only create report if category has documents
            print(f"Creating report for {category} with {len(docs)} documents...")
            category_filename = category.replace(" ", "_").lower()
            output_path = os.path.join(reports_dir, f"Alstom_{category_filename}_ollama.docx")
            create_category_report(category, docs, output_path)
    
    # Create master report
    master_path = os.path.join(reports_dir, "Alstom_Master_Index_ollama.docx")
    create_master_report(categorized_docs, master_path)
    
    print("\nAll reports generated successfully!")

if __name__ == "__main__":
    main()
