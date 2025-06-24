import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def format_number(num):
    """Format number with thousands separator"""
    return "{:,}".format(num)

def create_word_report(json_path, output_path):
    # Load the analysis data
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Create a new Word document
    doc = Document()
    
    # Add styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    doc.add_heading('Alstom Project Documentation Analysis', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run('This report provides a comprehensive analysis of the documentation available for the Alstom project. ')
    summary.add_run(f'A total of {format_number(data["total_files"])} files have been identified and categorized across various folders. ')
    summary.add_run('The documentation includes technical drawings, project correspondence, specifications, and other essential project documents.')
    
    # Key Statistics
    doc.add_heading('Key Statistics', level=1)
    
    # Group files by type
    drawings = data['extensions'].get('.dwg', 0)
    pdfs = data['extensions'].get('.pdf', 0)
    emails = data['extensions'].get('.msg', 0)
    word_docs = data['extensions'].get('.doc', 0) + data['extensions'].get('.docx', 0)
    excel_files = data['extensions'].get('.xls', 0) + data['extensions'].get('.xlsx', 0)
    images = data['extensions'].get('.jpg', 0) + data['extensions'].get('.jpeg', 0) + data['extensions'].get('.png', 0)
    
    # Create statistics table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Document Type'
    header_cells[1].text = 'Count'
    
    # Add data rows
    data_rows = [
        ('Technical Drawings (DWG)', drawings),
        ('PDF Documents', pdfs),
        ('Email Communications', emails),
        ('Word Documents', word_docs),
        ('Excel Spreadsheets', excel_files),
        ('Images', images),
        ('Total Files', data['total_files'])
    ]
    
    for doc_type, count in data_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = doc_type
        row_cells[1].text = format_number(count)
    
    doc.add_paragraph()
    
    # Document Types Analysis
    doc.add_heading('Document Types Analysis', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Technical Drawings: ').bold = True
    p.add_run(f'The project contains {format_number(drawings)} AutoCAD drawings (.dwg files), representing the technical specifications and design details.')
    
    p = doc.add_paragraph()
    p.add_run('PDF Documents: ').bold = True
    p.add_run(f'There are {format_number(pdfs)} PDF files, which may include approved drawings, reports, and other documentation.')
    
    p = doc.add_paragraph()
    p.add_run('Project Communications: ').bold = True
    p.add_run(f'The archive includes {format_number(emails)} email messages (.msg files), representing project correspondence and communications.')
    
    p = doc.add_paragraph()
    p.add_run('Office Documents: ').bold = True
    p.add_run(f'There are {format_number(word_docs)} Word documents and {format_number(excel_files)} Excel spreadsheets containing project specifications, reports, and data.')
    
    # Recommendations
    doc.add_heading('Recommendations', level=1)
    recommendations = [
        'Review the technical drawings to understand the scope and technical specifications of the project.',
        'Examine the PDF documents as they likely contain approved versions of important project documentation.',
        'Go through the email communications to understand project history and important decisions.',
        'Check Word documents for detailed specifications and reports.',
        'Review Excel files for any project data, calculations, or schedules.'
    ]
    
    for rec in recommendations:
        p = doc.add_paragraph(rec, style='List Bullet')
    
    # Next Steps
    doc.add_heading('Next Steps', level=1)
    p = doc.add_paragraph('To facilitate easier access to this documentation, we recommend:')
    next_steps = [
        'Setting up a document management system to organize these files',
        'Creating a detailed index of critical documents',
        'Identifying and prioritizing key documentation for review',
        'Establishing a system for tracking document revisions and updates'
    ]
    
    for step in next_steps:
        p = doc.add_paragraph(step, style='List Bullet')
    
    # Save the document
    doc.save(output_path)
    return output_path

def main():
    # Paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    json_path = os.path.join(script_dir, 'reports', 'file_analysis.json')
    output_path = os.path.join(script_dir, 'reports', 'Alstom_Project_Documentation_Analysis.docx')
    
    # Create reports directory if it doesn't exist
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Generate report
    print("Generating Word report...")
    report_path = create_word_report(json_path, output_path)
    print(f"\nReport generated successfully: {report_path}")

if __name__ == "__main__":
    main()
