import os
import base64
import subprocess
from PIL import Image
from pdf2image import convert_from_path
import pytesseract
from docx import Document
import time

PDF_SOURCE_FOLDER = r'H:\Projects Control (PC)\10 Backup\06 Yasser\Damietta Buildings Project'
OUTPUT_DIR = r'D:\OEssam\Test\phi3'
IMAGE_OUTPUT_DIR = os.path.join(OUTPUT_DIR, 'images')
PROCESSED_FILES_LOG = os.path.join(OUTPUT_DIR, 'processed_pdfs.txt')
LLM_MODEL_TAG = 'phi3:latest'
POPPLER_Path = r"C:\\Program Files\\poppler\\Library\\bin"
pytesseract.pytesseract.tesseract_cmd = r"C:\\Users\\Omar Essam2\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe"

PAGE_PROMPT = """You are an expert at reading scanned documents. Below is the OCR text and the image of a single page. Extract and summarize the essential information (project, contract value, parties, dates, key terms, obligations, etc). Be concise and clear.\n\nOCR Text:\n{text}\n"""

SUMMARY_PROMPT = """You are an expert summarizer. Here are the summaries of all pages in a scanned contract. Write a concise, comprehensive summary of the entire document, focusing on the most important information.\n\nPage Summaries:\n{text}\n\nFull Document Summary:"""

def get_processed_pdfs():
    if not os.path.exists(PROCESSED_FILES_LOG):
        return set()
    with open(PROCESSED_FILES_LOG, 'r', encoding='utf-8') as f:
        return set(line.strip() for line in f)

def mark_pdf_as_processed(pdf_path):
    with open(PROCESSED_FILES_LOG, 'a', encoding='utf-8') as f:
        f.write(pdf_path + '\n')

def get_all_unprocessed_pdfs():
    processed = get_processed_pdfs()
    pdfs = []
    for dirpath, _, filenames in os.walk(PDF_SOURCE_FOLDER):
        for fn in filenames:
            if fn.lower().endswith('.pdf'):
                full = os.path.join(dirpath, fn)
                if full not in processed:
                    pdfs.append(full)
    pdfs.sort()
    return pdfs

def pdf_to_images(pdf_path, image_save_dir):
    os.makedirs(image_save_dir, exist_ok=True)
    images = convert_from_path(pdf_path, dpi=300, poppler_path=POPPLER_Path)
    image_paths = []
    pdf_basename = os.path.splitext(os.path.basename(pdf_path))[0]
    for i, img in enumerate(images):
        out_name = f"{pdf_basename}_page{i+1}.png"
        path = os.path.join(image_save_dir, out_name)
        img.save(path, "PNG")
        image_paths.append(path)
    return image_paths

def run_tesseract(image_path):
    try:
        return pytesseract.image_to_string(Image.open(image_path), lang='ara+eng')
    except Exception as e:
        print(f"Tesseract error: {e}")
        return ""

def image_to_base64(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode("utf-8")

def query_phi3(prompt, image_b64=None):
    if image_b64:
        prompt = f"<|image|>{image_b64}\n{prompt}"
    start = time.time()
    process = subprocess.Popen([
        "ollama", "run", LLM_MODEL_TAG
    ], stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    stdout, stderr = process.communicate(prompt.encode("utf-8"))
    duration = time.time() - start
    out = stdout.decode("utf-8", errors="ignore")
    err = stderr.decode("utf-8", errors="ignore")
    if err:
        print(f"[phi3 ERROR]: {err}")
    return out.strip(), duration

def process_pdf(pdf_path):
    doc = Document()
    pdf_filename_base = os.path.splitext(os.path.basename(pdf_path))[0]
    doc.add_heading(f'Phi3 Multimodal Summary: {pdf_filename_base}', 0)
    doc.add_paragraph(f'Source PDF: {pdf_path}')
    print(f"\n=== Processing PDF: {pdf_path} ===")
    image_files = pdf_to_images(pdf_path, IMAGE_OUTPUT_DIR)
    all_page_summaries = []
    for idx, image_file_path in enumerate(image_files):
        page_num = idx + 1
        doc.add_heading(f'Page {page_num}', level=2)
        print(f"\n--- Processing Page {page_num} of {len(image_files)} ({os.path.basename(image_file_path)}) ---")
        ocr_text = run_tesseract(image_file_path)
        img_b64 = image_to_base64(image_file_path)
        page_prompt = PAGE_PROMPT.format(text=ocr_text)
        page_summary, duration = query_phi3(page_prompt, img_b64)
        all_page_summaries.append(page_summary)
        doc.add_paragraph(f'Duration: {duration:.2f} seconds')
        doc.add_paragraph('Raw OCR Preview:')
        doc.add_paragraph(ocr_text)
        doc.add_paragraph('Phi3 Page Summary:')
        doc.add_paragraph(page_summary)
        doc.add_paragraph('-' * 40)
    if all_page_summaries:
        print("\n--- Generating Full Document Summary --- ")
        concatenated = "\n\n".join(all_page_summaries)
        summary_prompt = SUMMARY_PROMPT.format(text=concatenated)
        final_summary, summary_duration = query_phi3(summary_prompt)
        doc.add_heading('Final Comprehensive Summary', level=1)
        doc.add_paragraph(f'Duration for summary generation: {summary_duration:.2f} seconds')
        doc.add_paragraph(final_summary)
    output_word_filename = f"phi3_multimodal_summary_{pdf_filename_base}.docx"
    result_docx_path = os.path.join(OUTPUT_DIR, output_word_filename)
    try:
        doc.save(result_docx_path)
        print(f"\nResults for {os.path.basename(pdf_path)} saved to {result_docx_path}")
        mark_pdf_as_processed(pdf_path)
    except Exception as e:
        print(f"[ERROR] Failed to save Word document {result_docx_path}: {e}")

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(IMAGE_OUTPUT_DIR, exist_ok=True)
    if not os.path.exists(PROCESSED_FILES_LOG):
        with open(PROCESSED_FILES_LOG, 'w', encoding='utf-8') as f:
            pass
    pdfs = get_all_unprocessed_pdfs()
    if not pdfs:
        print("No new PDFs to process.")
        return
    for pdf_path in pdfs:
        process_pdf(pdf_path)
        print("\n" + "="*50 + "\n")

if __name__ == "__main__":
    main()
